"""
YouTube 구독 채널 필터 검색 프로그램
- Eel 기반 데스크톱 앱
"""

import eel
import os
import sys
import json
import ctypes
import subprocess
import threading

# Windows 작업표시줄 아이콘 고정을 위한 AppUserModelID 설정
# 이 설정이 없으면 작업표시줄에 고정 시 아이콘이 변경됨
if sys.platform == 'win32':
    try:
        myappid = 'RoySearch.YouTubeSubscriptionSearch.1.0'
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
    except Exception:
        pass

from auth import (
    get_authenticated_service,
    is_configured, is_authenticated, logout,
    get_auth_url, exchange_code_for_token,
    get_auth_url_with_localhost, start_auth_server, open_auth_browser, find_free_port as find_auth_port
)
import account_manager
from youtube_api import get_subscriptions, get_channels_batch, get_videos_batch, get_channel_uploads, get_popular_videos, search_youtube_videos, get_filtered_comments
from rss_fetcher import fetch_all_channels
import cache_manager
import config
from data_path import (
    CONFIG_FILE, EXPORT_FILE, DATA_DIR, CHANNEL_FILE, CREDENTIALS_DIR,
    get_data_dir_path, get_oauth_json_dir, get_available_oauth_files,
    load_oauth_credentials, load_token_credentials, save_token_credentials,
    migrate_legacy_credentials, export_credentials_to_file, import_credentials_from_file
)
from secure_config import (
    save_secure_config, load_secure_config, has_secure_config,
    delete_secure_config, verify_password
)

# 전역 변수
youtube_service = None
subscriptions = []
search_cancelled = False
selected_channel_id = None  # 선택된 채널 ID

# 시작 시 마이그레이션 수행
account_manager.migrate_single_token()
# 레거시 자격증명 마이그레이션 (json 폴더 -> AppData)
migrate_legacy_credentials()

# RoyStudio 백엔드 모듈 로드
import studio_backend

# 채널 관리 모듈 로드
import channel_manager

# Eel 초기화 - exe에서는 내부 리소스 경로 사용
if getattr(sys, 'frozen', False):
    # PyInstaller로 빌드된 경우 _MEIPASS에서 web 폴더 찾기
    exe_dir = os.path.dirname(sys.executable)
    bundle_dir = getattr(sys, '_MEIPASS', exe_dir)
    eel.init(os.path.join(bundle_dir, 'web'))
else:
    eel.init('web')


@eel.expose
def get_config_status():
    """
    설정 상태를 반환합니다.

    각 계정은 자체 API 설정을 가지므로,
    계정이 있고 해당 계정에 API 설정이 되어 있어야 'isConfigured'가 True입니다.
    """
    global youtube_service

    # 현재 계정 확인
    current_account = account_manager.get_current_account()

    # 계정이 없으면 설정 필요 (첫 실행)
    if not current_account:
        return {
            'isConfigured': False,
            'isAuthenticated': False,
            'needsFirstSetup': True
        }

    # 계정의 API 설정이 있는지 확인
    has_api = account_manager.has_account_api_credentials(current_account['id'])

    if not has_api:
        return {
            'isConfigured': False,
            'isAuthenticated': False,
            'accountId': current_account['id'],
            'needsApiSetup': True
        }

    # API 설정이 있으면 자동으로 런타임에 로드
    if not is_configured():
        api_result = account_manager.load_account_api_credentials(current_account['id'])
        if api_result['success']:
            config.set_current_credentials(api_result['api_key'], api_result['client_id'], api_result['client_secret'])
        else:
            return {
                'isConfigured': False,
                'isAuthenticated': False,
                'accountId': current_account['id'],
                'needsApiSetup': True,
                'error': api_result.get('error', 'API 설정 로드 실패')
            }

    authenticated = is_authenticated(current_account['id'])

    # 인증된 상태라면 자동으로 서비스 연결 시도
    if authenticated and not youtube_service:
        try:
            youtube_service = get_authenticated_service(current_account['id'])
            # 서비스 생성 실패 시 인증 실패로 처리
            if not youtube_service:
                authenticated = False
        except Exception as e:
            print(f"자동 로그인 실패: {e}")
            authenticated = False

    return {
        'isConfigured': True,
        'isAuthenticated': authenticated,
        'accountId': current_account['id'],
        'accountName': current_account.get('name', '')
    }


@eel.expose
def get_preset_oauth_accounts():
    """
    사전 설정된 OAuth 계정 목록을 반환합니다.
    AppData와 레거시 json 폴더 모두에서 검색합니다.
    """
    try:
        oauth_files = get_available_oauth_files()
        print(f"[프리셋 계정] OAuth 파일 수: {len(oauth_files)}")
        for f in oauth_files:
            status = "암호화" if f.get('encrypted') else "레거시"
            token_status = "토큰있음" if f['hasToken'] else "토큰없음"
            print(f"  - {f['namePart']} ({status}, {token_status})")

        return {
            'success': True,
            'accounts': oauth_files,
            'hasPresetAccounts': len(oauth_files) > 0,
            'credentialsDir': CREDENTIALS_DIR
        }
    except Exception as e:
        print(f"OAuth 목록 조회 오류: {e}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': str(e),
            'accounts': [],
            'hasPresetAccounts': False
        }


# 현재 사용 중인 프리셋 OAuth 파일명 (토큰 저장용)
_current_preset_oauth_filename = None


# 현재 사용 중인 프리셋 계정 이름
_current_preset_name_part = None


@eel.expose
def login(client_id, client_secret):
    """
    Client ID와 Client Secret으로 직접 로그인합니다.
    OAuth 인증 후 토큰을 생성합니다.
    """
    global youtube_service

    try:
        from google_auth_oauthlib.flow import InstalledAppFlow
        from googleapiclient.discovery import build

        # OAuth 설정
        config.set_current_credentials('', client_id, client_secret)

        # OAuth 데이터 구성
        oauth_data = {
            'installed': {
                'client_id': client_id,
                'client_secret': client_secret,
                'auth_uri': 'https://accounts.google.com/o/oauth2/auth',
                'token_uri': 'https://oauth2.googleapis.com/token',
                'redirect_uris': ['http://localhost']
            }
        }

        SCOPES = [
            'https://www.googleapis.com/auth/youtube',
            'https://www.googleapis.com/auth/youtube.force-ssl'
        ]

        print("\n[로그인] 브라우저에서 구글 로그인을 진행하세요.")

        flow = InstalledAppFlow.from_client_config(oauth_data, SCOPES)
        creds = flow.run_local_server(
            port=0,
            prompt='consent',
            success_message='인증 완료! 이 창을 닫아도 됩니다.',
            open_browser=True
        )

        # YouTube 서비스 생성
        youtube_service = build('youtube', 'v3', credentials=creds)

        print("[로그인] 로그인 성공!")
        return {'success': True}

    except Exception as e:
        print(f"[로그인] 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def login_with_preset_oauth(name_part, auto_login=False):
    """
    사전 설정된 OAuth 자격증명을 사용하여 로그인합니다.
    암호화된 파일과 레거시 파일 모두 지원합니다.

    Args:
        name_part: 계정 이름 (예: '득수_getwater')
        auto_login: True면 토큰이 있을 때 자동 로그인 시도
    """
    global youtube_service, _current_preset_name_part, subscriptions

    try:
        # OAuth 자격증명 로드 (암호화/레거시 자동 감지)
        oauth_data = load_oauth_credentials(name_part)

        if not oauth_data:
            return {'success': False, 'error': f'OAuth 자격증명을 찾을 수 없습니다: {name_part}'}

        # installed 형식 확인
        if 'installed' not in oauth_data:
            return {'success': False, 'error': 'OAuth 자격증명 형식이 올바르지 않습니다.'}

        installed = oauth_data['installed']
        client_id = installed.get('client_id')
        client_secret = installed.get('client_secret')

        if not client_id or not client_secret:
            return {'success': False, 'error': 'OAuth 자격증명에 client_id 또는 client_secret이 없습니다.'}

        # 런타임 config에 적용
        config.set_current_credentials('', client_id, client_secret)

        # 현재 프리셋 계정 이름 저장
        _current_preset_name_part = name_part

        # 캐시 매니저에 프리셋 계정 ID 설정 (계정별 캐시 분리)
        cache_manager.set_current_preset_account(name_part)

        # 계정 전환 시 기존 구독 목록 초기화
        subscriptions = []

        # 토큰 로드 시도
        if auto_login:
            token_data = load_token_credentials(name_part)

            if token_data:
                try:
                    from google.oauth2.credentials import Credentials
                    from google.auth.transport.requests import Request
                    from googleapiclient.discovery import build

                    creds = Credentials.from_authorized_user_info(token_data)

                    # 토큰 만료 시 갱신
                    if creds and creds.expired and creds.refresh_token:
                        creds.refresh(Request())
                        # 갱신된 토큰 저장 (암호화)
                        save_token_credentials(name_part, creds.to_json())

                    if creds and creds.valid:
                        youtube_service = build('youtube', 'v3', credentials=creds)
                        return {
                            'success': True,
                            'autoLogin': True,
                            'message': f'{name_part} 계정으로 자동 로그인되었습니다.',
                            'accountName': name_part
                        }
                except Exception as e:
                    print(f"자동 로그인 실패, 수동 로그인 필요: {e}")
                    # 토큰이 유효하지 않으면 삭제
                    from data_path import delete_token_credentials
                    delete_token_credentials(name_part)

        return {
            'success': True,
            'autoLogin': False,
            'message': f'{name_part} 계정으로 설정되었습니다. 로그인을 진행해주세요.',
            'accountName': name_part,
            'needsLogin': True
        }

    except Exception as e:
        print(f"OAuth 로그인 설정 오류: {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def save_api_config(client_id, client_secret):
    """
    [DEPRECATED] 전역 API 설정 - 더 이상 사용하지 않음.
    계정별 API 설정(save_account_api_config)을 사용하세요.
    """
    return {'success': False, 'error': '전역 API 설정은 더 이상 지원하지 않습니다. 계정별 API 설정을 사용하세요.'}


# ===== 토큰 생성 API (토큰생성기 기능 통합) =====

@eel.expose
def create_token_for_account(name_part):
    """
    특정 계정의 토큰을 생성합니다.
    브라우저 인증을 통해 토큰을 생성하고 암호화하여 저장합니다.

    Args:
        name_part: 계정 이름 (예: '득수_getwater')

    Returns:
        dict: {'success': bool, 'error': str}
    """
    global youtube_service, _current_preset_name_part

    try:
        from google_auth_oauthlib.flow import InstalledAppFlow
        from googleapiclient.discovery import build

        # OAuth 자격증명 로드
        oauth_data = load_oauth_credentials(name_part)
        if not oauth_data:
            return {'success': False, 'error': f'OAuth 자격증명을 찾을 수 없습니다: {name_part}'}

        # OAuth 스코프
        SCOPES = [
            'https://www.googleapis.com/auth/youtube',
            'https://www.googleapis.com/auth/youtube.force-ssl'
        ]

        print(f"\n[토큰 생성] {name_part} 계정의 토큰 생성 중...")
        print("브라우저에서 구글 로그인을 진행하세요.")

        # Flow 생성 및 인증
        flow = InstalledAppFlow.from_client_config(oauth_data, SCOPES)

        # 로컬 서버로 인증 (브라우저 자동 열림)
        creds = flow.run_local_server(
            port=0,  # 사용 가능한 포트 자동 선택
            prompt='consent',
            success_message='인증 완료! 이 창을 닫아도 됩니다.',
            open_browser=True
        )

        # 토큰 암호화 저장
        token_json = creds.to_json()
        if save_token_credentials(name_part, token_json):
            print(f"[토큰 생성] 토큰 저장 완료: {name_part}")

            # 현재 계정으로 설정하고 YouTube 서비스 생성
            _current_preset_name_part = name_part

            # config 설정
            installed = oauth_data.get('installed', {})
            config.set_current_credentials('', installed.get('client_id', ''), installed.get('client_secret', ''))

            # 캐시 매니저 설정
            cache_manager.set_current_preset_account(name_part)

            # YouTube 서비스 생성
            youtube_service = build('youtube', 'v3', credentials=creds)

            return {
                'success': True,
                'message': f'{name_part} 토큰 생성 및 로그인 완료!',
                'accountName': name_part
            }
        else:
            return {'success': False, 'error': '토큰 저장 실패'}

    except Exception as e:
        print(f"[토큰 생성] 오류: {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def delete_account_token(name_part):
    """
    특정 계정의 토큰을 삭제합니다.

    Args:
        name_part: 계정 이름

    Returns:
        dict: {'success': bool}
    """
    try:
        from data_path import delete_token_credentials
        delete_token_credentials(name_part)
        return {'success': True, 'message': f'{name_part} 토큰이 삭제되었습니다.'}
    except Exception as e:
        return {'success': False, 'error': str(e)}


# ===== 자격증명 내보내기/가져오기 API =====

@eel.expose
def export_all_credentials(password):
    """
    모든 자격증명을 암호화하여 내보냅니다.
    사용자가 파일 저장 위치를 선택합니다.

    Args:
        password: 암호화 비밀번호

    Returns:
        dict: {'success': bool, 'count': int, 'path': str, 'error': str}
    """
    try:
        import tkinter as tk
        from tkinter import filedialog

        # 파일 저장 대화상자
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        file_path = filedialog.asksaveasfilename(
            title='자격증명 내보내기',
            defaultextension='.roycred',
            filetypes=[('로이 자격증명 파일', '*.roycred'), ('모든 파일', '*.*')],
            initialfilename='credentials_backup.roycred'
        )

        root.destroy()

        if not file_path:
            return {'success': False, 'error': '파일 저장이 취소되었습니다.'}

        result = export_credentials_to_file(file_path, password)

        if result['success']:
            return {
                'success': True,
                'count': result['count'],
                'path': file_path,
                'message': f'{result["count"]}개 계정의 자격증명을 내보냈습니다.'
            }
        else:
            return result

    except Exception as e:
        return {'success': False, 'error': str(e)}


@eel.expose
def import_all_credentials(password):
    """
    파일에서 자격증명을 가져옵니다.

    Args:
        password: 복호화 비밀번호

    Returns:
        dict: {'success': bool, 'count': int, 'error': str}
    """
    try:
        import tkinter as tk
        from tkinter import filedialog

        # 파일 선택 대화상자
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        file_path = filedialog.askopenfilename(
            title='자격증명 가져오기',
            filetypes=[('로이 자격증명 파일', '*.roycred'), ('모든 파일', '*.*')]
        )

        root.destroy()

        if not file_path:
            return {'success': False, 'error': '파일 선택이 취소되었습니다.'}

        result = import_credentials_from_file(file_path, password)

        if result['success']:
            return {
                'success': True,
                'count': result['count'],
                'message': f'{result["count"]}개 계정의 자격증명을 가져왔습니다.'
            }
        else:
            return result

    except Exception as e:
        return {'success': False, 'error': str(e)}


@eel.expose
def add_oauth_account_from_file():
    """
    OAuth JSON 파일을 선택하여 새 계정을 추가합니다.
    파일을 암호화하여 AppData에 저장합니다.

    Returns:
        dict: {'success': bool, 'name_part': str, 'error': str}
    """
    try:
        import tkinter as tk
        from tkinter import filedialog

        # 파일 선택 대화상자
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        file_path = filedialog.askopenfilename(
            title='OAuth JSON 파일 선택',
            filetypes=[('OAuth JSON 파일', '*_OAuth.json'), ('JSON 파일', '*.json'), ('모든 파일', '*.*')]
        )

        root.destroy()

        if not file_path:
            return {'success': False, 'error': '파일 선택이 취소되었습니다.'}

        # 파일명 추출
        filename = os.path.basename(file_path)

        # 파일명 형식 검증 (xxx_OAuth.json 또는 xxx_yyy_OAuth.json)
        if not filename.endswith('.json'):
            return {'success': False, 'error': 'JSON 파일을 선택해주세요.'}

        # name_part 추출 (파일명에서 _OAuth.json 제거)
        if filename.endswith('_OAuth.json'):
            name_part = filename.replace('_OAuth.json', '')
        else:
            # _OAuth.json 형식이 아니면 .json만 제거하고 _OAuth 없이 사용
            name_part = filename.replace('.json', '')

        # 이미 존재하는 계정인지 확인
        existing_files = get_available_oauth_files()
        for existing in existing_files:
            if existing['namePart'] == name_part:
                return {'success': False, 'error': f'이미 등록된 계정입니다: {name_part}'}

        # OAuth JSON 파일 읽기
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                oauth_data = json.load(f)
        except json.JSONDecodeError:
            return {'success': False, 'error': 'JSON 파일 형식이 올바르지 않습니다.'}

        # OAuth 데이터 유효성 검증 (installed 또는 web 키가 있어야 함)
        if 'installed' not in oauth_data and 'web' not in oauth_data:
            return {'success': False, 'error': 'OAuth 자격증명 파일이 아닙니다. installed 또는 web 키가 없습니다.'}

        # 암호화하여 저장
        from data_path import encrypt_credentials, ensure_credentials_dir
        ensure_credentials_dir()

        encrypted = encrypt_credentials(oauth_data)
        if not encrypted:
            return {'success': False, 'error': '암호화 실패'}

        encrypted_path = os.path.join(CREDENTIALS_DIR, f'{name_part}_OAuth.enc')
        with open(encrypted_path, 'wb') as f:
            f.write(encrypted)

        # 표시 이름 추출
        parts = name_part.split('_')
        display_name = parts[0] if parts else name_part
        email_hint = parts[1] if len(parts) >= 2 else ''

        print(f"[계정 추가] 새 계정 추가됨: {name_part}")

        return {
            'success': True,
            'name_part': name_part,
            'display_name': display_name,
            'email': email_hint,
            'message': f'계정이 추가되었습니다: {display_name}'
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def get_credentials_dir():
    """자격증명 저장 경로를 반환합니다."""
    return {
        'path': CREDENTIALS_DIR,
        'exists': os.path.exists(CREDENTIALS_DIR)
    }


@eel.expose
def get_api_config():
    """
    현재 런타임 API 설정 상태를 반환합니다.
    계정별 API 시스템에서는 런타임에 로드된 자격 증명을 확인합니다.
    """
    client_id = config.get_client_id() or ''
    client_secret = config.get_client_secret() or ''

    # 현재 계정의 API 상태 확인
    current_account = account_manager.get_current_account()
    has_account_api = False
    if current_account:
        has_account_api = account_manager.has_account_api_credentials(current_account['id'])

    return {
        'clientId': client_id[:20] + '...' if client_id else '',  # 보안상 일부만 표시
        'clientSecret': '***' if client_secret else '',
        'hasConfig': bool(client_id and client_secret),
        'hasAccountApi': has_account_api
    }


@eel.expose
def save_credentials(client_id, client_secret, password):
    """
    [DEPRECATED] 전역 암호화 저장 - 더 이상 사용하지 않음.
    계정별 API 설정(save_account_api_config)을 사용하세요.
    """
    return {'success': False, 'error': '전역 API 설정은 더 이상 지원하지 않습니다. 계정별 API 설정을 사용하세요.'}


@eel.expose
def load_credentials():
    """
    현재 계정의 API 자격 증명을 런타임에 로드합니다.
    """
    current_account = account_manager.get_current_account()
    if not current_account:
        return {'success': False, 'error': '현재 계정이 없습니다.'}

    result = account_manager.load_account_api_credentials(current_account['id'])
    if result['success']:
        # 런타임 config 모듈에 적용
        config.set_current_credentials(result['api_key'], result['client_id'], result['client_secret'])
        return {'success': True}
    return result


@eel.expose
def check_saved_credentials():
    """현재 계정에 저장된 API 인증 정보가 있는지 확인합니다."""
    current_account = account_manager.get_current_account()
    if not current_account:
        return {'hasSavedCredentials': False, 'hasAccount': False}

    has_api = account_manager.has_account_api_credentials(current_account['id'])
    return {
        'hasSavedCredentials': has_api,
        'hasAccount': True,
        'accountId': current_account['id'],
        'accountName': current_account.get('name', '')
    }


@eel.expose
def delete_saved_credentials():
    """현재 계정의 저장된 API 인증 정보를 삭제합니다."""
    current_account = account_manager.get_current_account()
    if not current_account:
        return {'success': False, 'error': '현재 계정이 없습니다.'}

    success = account_manager.delete_account_api_credentials(current_account['id'])
    if success:
        config.clear_current_credentials()
    return {'success': success}


# 인증 플로우 저장 (코드 입력 대기용)
_pending_auth_flow = None


@eel.expose
def start_login():
    """
    OAuth 로그인을 시작합니다. (수동 코드 입력용 - fallback)
    브라우저에서 열 URL을 반환합니다.
    """
    global _pending_auth_flow

    if not is_configured():
        return {'success': False, 'error': 'API 설정이 필요합니다.'}

    try:
        flow, auth_url = get_auth_url()
        if not flow or not auth_url:
            return {'success': False, 'error': 'API 설정을 확인해주세요.'}

        # 플로우 저장 (코드 입력 시 사용)
        _pending_auth_flow = flow

        return {
            'success': True,
            'authUrl': auth_url
        }
    except Exception as e:
        print(f"인증 URL 생성 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def start_login_with_browser():
    """
    전용 Chrome 창에서 OAuth 로그인을 시작합니다.
    localhost 리다이렉트로 자동으로 코드를 받습니다.
    프리셋 OAuth 사용 시 토큰을 json 폴더에 저장합니다.
    """
    global youtube_service, _current_preset_oauth_filename
    import threading
    from data_path import get_token_path_for_oauth

    if not is_configured():
        return {'success': False, 'error': 'API 설정이 필요합니다.'}

    try:
        # 사용 가능한 포트 찾기
        port = find_auth_port(8888)
        print(f"인증 서버 포트: {port}")

        # 인증 URL 생성
        flow, auth_url = get_auth_url_with_localhost(port)
        if not flow or not auth_url:
            return {'success': False, 'error': 'API 설정을 확인해주세요.'}

        # 전용 Chrome 창에서 인증 URL 열기
        # 프리셋 OAuth 사용 시 계정별 별도 Chrome 프로필 사용
        print("전용 Chrome 창에서 로그인 페이지를 엽니다...")
        profile_name = None
        if _current_preset_oauth_filename:
            # OAuth 파일명에서 프로필 이름 추출 (예: 로이_roy_OAuth.json -> 로이_roy)
            profile_name = _current_preset_oauth_filename.replace('_OAuth.json', '')
        open_auth_browser(auth_url, profile_name)

        # 인증 콜백 서버 시작 (블로킹)
        print("인증 대기 중...")
        auth_result = start_auth_server(port)

        if auth_result.get('error'):
            return {'success': False, 'error': f"로그인 실패: {auth_result['error']}"}

        if not auth_result.get('code'):
            return {'success': False, 'error': '인증 코드를 받지 못했습니다.'}

        # 코드로 토큰 교환
        print("토큰 교환 중...")
        creds = exchange_code_for_token(flow, auth_result['code'])

        if not creds:
            return {'success': False, 'error': '토큰 교환에 실패했습니다.'}

        # 프리셋 OAuth 사용 중이면 토큰을 json 폴더에 저장
        if _current_preset_oauth_filename:
            try:
                token_path = get_token_path_for_oauth(_current_preset_oauth_filename)
                with open(token_path, 'w', encoding='utf-8') as f:
                    f.write(creds.to_json())
                print(f"토큰 저장됨: {token_path}")
            except Exception as e:
                print(f"토큰 저장 실패 (무시됨): {e}")

        # YouTube 서비스 생성
        from googleapiclient.discovery import build
        youtube_service = build('youtube', 'v3', credentials=creds)

        print("로그인 성공!")
        return {'success': True}

    except Exception as e:
        print(f"로그인 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def complete_login(auth_code):
    """
    OAuth 로그인을 완료합니다. (2단계: 인증 코드로 토큰 교환)
    """
    global youtube_service, _pending_auth_flow

    if not _pending_auth_flow:
        return {'success': False, 'error': '로그인을 다시 시작해주세요.'}

    if not auth_code or not auth_code.strip():
        return {'success': False, 'error': '인증 코드를 입력해주세요.'}

    try:
        creds = exchange_code_for_token(_pending_auth_flow, auth_code.strip())
        _pending_auth_flow = None  # 플로우 초기화

        if not creds:
            return {'success': False, 'error': '인증 코드가 올바르지 않습니다.\n다시 시도해주세요.'}

        # YouTube 서비스 생성
        from googleapiclient.discovery import build
        youtube_service = build('youtube', 'v3', credentials=creds)

        return {'success': True}
    except Exception as e:
        print(f"로그인 완료 오류: {e}")
        _pending_auth_flow = None
        return {'success': False, 'error': f'로그인 실패: {str(e)}'}


@eel.expose
def do_login():
    """OAuth 로그인을 수행합니다. (기존 토큰이 있는 경우 자동 로그인)"""
    global youtube_service

    if not is_configured():
        return {'success': False, 'error': 'API 설정이 필요합니다.'}

    try:
        youtube_service = get_authenticated_service()
        if youtube_service:
            return {'success': True}
        # 새 로그인이 필요한 경우 - UI에서 start_login 호출 필요
        return {'success': False, 'needsManualLogin': True, 'error': '로그인이 필요합니다.'}
    except Exception as e:
        print(f"로그인 오류: {e}")
        return {'success': False, 'error': '로그인 중 오류가 발생했습니다.\n다시 시도해주세요.'}


@eel.expose
def get_user_channels():
    """현재 로그인한 사용자의 모든 채널 목록을 반환합니다."""
    global youtube_service, selected_channel_id

    try:
        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service:
            return {'success': False, 'channels': [], 'selectedChannelId': None}

        # 내 채널 정보 조회 (managedByMe=True는 브랜드 채널 포함)
        channels = []

        # 1. 기본 채널 조회
        request = youtube_service.channels().list(
            part='snippet,contentDetails',
            mine=True
        )
        response = request.execute()

        for item in response.get('items', []):
            channels.append({
                'id': item['id'],
                'title': item['snippet']['title'],
                'thumbnail': item['snippet']['thumbnails']['default']['url'],
                'isDefault': True
            })

        # 2. 관리하는 채널 조회 (브랜드 채널)
        try:
            request = youtube_service.channels().list(
                part='snippet,contentDetails',
                managedByMe=True,
                maxResults=50
            )
            response = request.execute()

            for item in response.get('items', []):
                # 이미 추가된 채널이 아니면 추가
                if not any(c['id'] == item['id'] for c in channels):
                    channels.append({
                        'id': item['id'],
                        'title': item['snippet']['title'],
                        'thumbnail': item['snippet']['thumbnails']['default']['url'],
                        'isDefault': False
                    })
        except Exception:
            # 브랜드 채널 조회는 MCN/콘텐츠 소유자만 가능 (일반 사용자는 403 오류)
            pass

        # 저장된 채널 ID 로드
        saved_channel_id = load_selected_channel()
        if saved_channel_id and any(c['id'] == saved_channel_id for c in channels):
            selected_channel_id = saved_channel_id
        elif channels:
            selected_channel_id = channels[0]['id']

        return {
            'success': True,
            'channels': channels,
            'selectedChannelId': selected_channel_id
        }
    except Exception as e:
        print(f"채널 목록 조회 오류: {e}")
        return {'success': False, 'channels': [], 'selectedChannelId': None}


@eel.expose
def select_channel(channel_id):
    """채널을 선택하고 저장합니다."""
    global selected_channel_id, subscriptions

    selected_channel_id = channel_id
    save_selected_channel(channel_id)

    # 채널 변경 시 구독 목록 캐시 초기화
    subscriptions = []
    cache_manager.clear_all_cache()

    return {'success': True}


@eel.expose
def get_current_channel():
    """현재 선택된 채널 정보를 반환합니다."""
    global youtube_service, selected_channel_id

    try:
        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service or not selected_channel_id:
            return {'success': False, 'channel': None}

        # 선택된 채널 정보 조회
        request = youtube_service.channels().list(
            part='snippet',
            id=selected_channel_id
        )
        response = request.execute()

        if response.get('items'):
            item = response['items'][0]
            return {
                'success': True,
                'channel': {
                    'id': item['id'],
                    'title': item['snippet']['title'],
                    'thumbnail': item['snippet']['thumbnails']['default']['url']
                }
            }

        return {'success': False, 'channel': None}
    except Exception as e:
        print(f"채널 정보 조회 오류: {e}")
        return {'success': False, 'channel': None}


def save_selected_channel(channel_id):
    """선택된 채널 ID를 파일에 저장합니다."""
    try:
        with open(CHANNEL_FILE, 'w', encoding='utf-8') as f:
            json.dump({'channelId': channel_id}, f)
    except Exception as e:
        print(f"채널 저장 실패: {e}")


def load_selected_channel():
    """저장된 채널 ID를 불러옵니다."""
    try:
        if os.path.exists(CHANNEL_FILE):
            with open(CHANNEL_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return data.get('channelId')
    except Exception as e:
        print(f"채널 로드 실패: {e}")
    return None


@eel.expose
def get_user_email():
    """현재 선택된 채널의 이름을 반환합니다."""
    global youtube_service, selected_channel_id

    try:
        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service:
            return {'success': False, 'email': None}

        # 선택된 채널이 있으면 해당 채널 정보 조회
        if selected_channel_id:
            request = youtube_service.channels().list(
                part='snippet',
                id=selected_channel_id
            )
        else:
            # 없으면 기본 채널
            request = youtube_service.channels().list(
                part='snippet',
                mine=True
            )

        response = request.execute()

        if response.get('items'):
            channel_title = response['items'][0]['snippet']['title']
            return {'success': True, 'email': channel_title}

        return {'success': False, 'email': None}
    except Exception as e:
        print(f"사용자 정보 조회 오류: {e}")
        return {'success': False, 'email': None}


@eel.expose
def do_logout():
    """로그아웃하고 프로그램을 종료합니다."""
    global youtube_service, subscriptions

    logout()
    cache_manager.clear_all_cache()
    youtube_service = None
    subscriptions = []

    # 프로그램 종료
    print("로그아웃 완료. 프로그램을 종료합니다.")
    os._exit(0)


@eel.expose
def load_subscriptions(force_refresh=False):
    """
    구독 채널 목록을 불러옵니다.
    """
    global youtube_service, subscriptions

    print(f"load_subscriptions 호출됨 - force_refresh: {force_refresh}")
    current_account = account_manager.get_current_account()
    print(f"현재 계정: {current_account['name'] if current_account else 'None'}")

    # 캐시 확인
    if not force_refresh:
        cached = cache_manager.load_subscriptions()
        if cached:
            print(f"캐시에서 {len(cached)}개 구독 채널 로드됨")
            # 캐시에 구독자 수가 없으면 API로 조회
            needs_subscriber_count = any(
                'subscriberCount' not in sub or sub.get('subscriberCount') == 0
                for sub in cached
            )

            if needs_subscriber_count:
                try:
                    if not youtube_service:
                        youtube_service = get_authenticated_service()

                    if youtube_service:
                        from youtube_api import get_channels_batch
                        channel_ids = [sub['id'] for sub in cached]
                        channel_stats = get_channels_batch(youtube_service, channel_ids)

                        for sub in cached:
                            stats = channel_stats.get(sub['id'], {})
                            sub['subscriberCount'] = stats.get('subscriberCount', 0)

                        cache_manager.save_subscriptions(cached)
                except Exception as e:
                    print(f"구독자 수 조회 실패: {e}")

            subscriptions = cached
            return {
                'success': True,
                'subscriptions': cached,
                'fromCache': True
            }

    # API 호출
    try:
        print("API 호출로 구독 채널 목록 가져오기...")
        if not youtube_service:
            print("youtube_service가 없음, get_authenticated_service() 호출")
            youtube_service = get_authenticated_service()

        if not youtube_service:
            print("youtube_service 생성 실패")
            return {'success': False, 'error': '로그인이 필요합니다.'}

        print("구독 채널 목록을 가져오는 중...")
        subs = get_subscriptions(youtube_service)

        if not subs:
            print("구독 채널이 없습니다.")
            return {'success': False, 'error': '구독 채널이 없습니다.'}

        print(f"API에서 {len(subs)}개 구독 채널 가져옴")
        cache_manager.save_subscriptions(subs)
        subscriptions = subs

        return {
            'success': True,
            'subscriptions': subs,
            'fromCache': False
        }

    except Exception as e:
        print(f"load_subscriptions 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def search_popular_videos(region_code='KR', category='0'):
    """
    국가별 인기 동영상을 검색합니다.

    Args:
        region_code: 국가 코드 (기본: KR)
        category: 카테고리 ('0'=전체, 또는 카테고리 ID)

    Returns:
        dict: {'success': True, 'videos': [...], 'stats': {...}}
    """
    global youtube_service

    try:
        # OAuth 서비스 사용
        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service:
            return {'success': False, 'error': '로그인이 필요합니다.'}

        print(f"인기 동영상 검색 시작... (국가: {region_code}, 카테고리: {category})")

        eel.update_progress("인기 동영상 조회 중...", 30)()

        all_videos = []

        if category == '0':
            # 전체 조회 (카테고리 필터 없음)
            all_videos = get_popular_videos(youtube_service, region_code=region_code, max_results=50)
        else:
            # 특정 카테고리만 조회
            try:
                all_videos = get_popular_videos(
                    youtube_service,
                    region_code=region_code,
                    video_category_id=category,
                    max_results=50
                )
            except Exception as e:
                # 해당 국가에서 카테고리가 지원되지 않으면 전체 조회로 폴백
                print(f"카테고리 {category} 조회 실패, 전체 조회로 대체: {e}")
                all_videos = get_popular_videos(youtube_service, region_code=region_code, max_results=50)

        eel.update_progress("완료!", 100)()

        # 조회수 내림차순 정렬
        all_videos.sort(key=lambda x: x.get('viewCount', 0), reverse=True)

        print(f"인기 동영상 {len(all_videos)}개 조회됨")

        return {
            'success': True,
            'videos': all_videos,
            'stats': {
                'total': len(all_videos),
                'filtered': len(all_videos),
                'regionCode': region_code
            }
        }

    except Exception as e:
        print(f"인기 동영상 검색 실패: {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def search_youtube_global(keyword, days_within=7, video_type='long'):
    """
    YouTube 전체에서 키워드로 영상을 검색합니다.
    """
    global youtube_service

    if not keyword:
        return {'success': False, 'error': '검색어를 입력하세요.'}

    # days_within을 정수로 변환
    days_within = int(days_within)

    try:
        # OAuth 서비스 사용
        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service:
            return {'success': False, 'error': '로그인이 필요합니다.'}

        print(f"YouTube 전체 검색: '{keyword}' (기간: {days_within}일, 타입: {video_type})")
        eel.update_progress("YouTube 검색 중...", 30)()

        # YouTube 검색 API 호출
        videos = search_youtube_videos(
            youtube_service,
            query=keyword,
            days_within=days_within,
            video_type=video_type,
            max_results=50
        )

        eel.update_progress("완료!", 100)()

        # 조회수 내림차순 정렬
        videos.sort(key=lambda x: x.get('viewCount', 0), reverse=True)

        print(f"YouTube 검색 결과: {len(videos)}개 영상")

        return {
            'success': True,
            'videos': videos,
            'stats': {
                'total': len(videos),
                'filtered': len(videos),
                'keyword': keyword
            }
        }

    except Exception as e:
        print(f"YouTube 검색 실패: {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def cancel_search():
    """검색을 중단합니다."""
    global search_cancelled
    search_cancelled = True
    print("검색 중단 요청됨")
    return {'success': True}


@eel.expose
def search_videos(filter_config):
    """
    조건에 맞는 영상을 검색합니다.
    """
    global youtube_service, subscriptions, search_cancelled

    # 검색 시작 시 취소 플래그 초기화
    search_cancelled = False

    if not subscriptions:
        return {'success': False, 'error': '먼저 구독 채널을 불러오세요.'}

    try:
        # OAuth 서비스 사용
        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service:
            return {'success': False, 'error': '로그인이 필요합니다.'}

        filter_type = filter_config.get('filterType', 'channel-monitor')
        video_type = filter_config.get('videoType', 'long')  # 'long' 또는 'shorts'
        max_subscribers = filter_config.get('maxSubscribers', 10000)
        min_views = filter_config.get('minViews', 10000)
        days_within_raw = filter_config.get('daysWithin', 15)
        mutation_ratio = filter_config.get('mutationRatio', 1.0)
        keyword = filter_config.get('keyword', '')

        # RSS 전용 모드 확인
        rss_only_mode = days_within_raw == 'rss'
        days_within = 15 if rss_only_mode else int(days_within_raw)  # RSS 모드는 15일 기본값 사용

        # 채널 ID 필터링: 카테고리 선택 시 해당 채널만, 그 외에는 전체
        filter_channel_ids = filter_config.get('channelIds')
        if filter_channel_ids:
            channel_ids = filter_channel_ids
        else:
            channel_ids = [sub['id'] for sub in subscriptions]
        print(f"총 {len(channel_ids)}개 채널 검색 시작... (필터: {filter_type})")

        # 취소 확인
        if search_cancelled:
            return {'success': False, 'error': '검색이 중단되었습니다.', 'cancelled': True}

        # RSS 전용 모드 여부에 따라 처리
        if rss_only_mode:
            print(f"RSS 전용 모드: {len(channel_ids)}개 채널")
            # 캐시된 채널 정보 사용 (구독 목록에서 가져온 정보)
            channel_info = {}
            for sub in subscriptions:
                channel_info[sub['id']] = {
                    'subscriberCount': sub.get('subscriberCount', 0),
                    'title': sub.get('title', ''),
                    'thumbnail': sub.get('thumbnail', '')
                }
        else:
            # 1단계: 채널 구독자 수 조회
            print("1단계: 채널 정보 조회 중...")
            eel.update_progress("채널 정보 조회 중...", 10)()
            channel_info = get_channels_batch(youtube_service, channel_ids)

        # 취소 확인
        if search_cancelled:
            return {'success': False, 'error': '검색이 중단되었습니다.', 'cancelled': True}

        # 2단계: RSS로 최신 영상 수집 (채널당 최대 15개)
        print("2단계: RSS 피드 수집 중...")
        eel.update_progress("RSS 피드 수집 중...", 20)()

        def rss_progress(current, total):
            percent = 20 + int((current / total) * 30)
            eel.update_progress(f"RSS 수집: {current}/{total}", percent)()
            # RSS 수집 중에도 취소 확인
            return not search_cancelled

        all_videos = fetch_all_channels(channel_ids, days_within, rss_progress)
        rss_video_count = len(all_videos)
        print(f"RSS에서 {rss_video_count}개 영상 수집됨")

        # 취소 확인
        if search_cancelled:
            return {'success': False, 'error': '검색이 중단되었습니다.', 'cancelled': True}

        # 2.5단계: 하이브리드 - RSS로 15개가 모두 채워진 채널은 API로 추가 조회
        # (기간이 15일 이상이거나 영상이 많은 채널의 경우)
        # RSS 전용 모드에서는 건너뜀
        if not rss_only_mode and days_within > 7:  # 7일 초과 기간일 때만 하이브리드 적용
            print("2.5단계: API로 추가 영상 조회 중...")
            eel.update_progress("API로 추가 조회 중...", 55)()

            # RSS에서 채널별 영상 수 계산
            channel_video_counts = {}
            for video in all_videos:
                cid = video['channelId']
                channel_video_counts[cid] = channel_video_counts.get(cid, 0) + 1

            # RSS에서 15개 영상이 수집된 채널 (더 있을 가능성)
            channels_need_api = [
                cid for cid, count in channel_video_counts.items()
                if count >= 15
            ]

            if channels_need_api:
                print(f"  - {len(channels_need_api)}개 채널에서 추가 조회 필요")

                # 기존 비디오 ID 집합 (중복 방지용)
                existing_video_ids = {v['videoId'] for v in all_videos}

                api_video_count = 0
                for i, cid in enumerate(channels_need_api):
                    if search_cancelled:
                        break

                    # 진행률 업데이트
                    percent = 55 + int((i / len(channels_need_api)) * 15)
                    eel.update_progress(f"API 조회: {i+1}/{len(channels_need_api)}", percent)()

                    # playlistItems API로 추가 영상 조회 (최대 50개, 기간 내)
                    api_videos = get_channel_uploads(
                        youtube_service,
                        cid,
                        days_within=days_within,
                        max_results=50
                    )

                    # 중복 제거 후 추가
                    for video in api_videos:
                        if video['videoId'] not in existing_video_ids:
                            all_videos.append(video)
                            existing_video_ids.add(video['videoId'])
                            api_video_count += 1

                print(f"  - API에서 {api_video_count}개 영상 추가됨")

        print(f"총 {len(all_videos)}개 영상 수집됨 (RSS: {rss_video_count}, API: {len(all_videos) - rss_video_count})")

        # 취소 확인
        if search_cancelled:
            return {'success': False, 'error': '검색이 중단되었습니다.', 'cancelled': True}

        if not all_videos:
            return {
                'success': True,
                'videos': [],
                'stats': {'total': 0, 'filtered': 0, 'rssMode': rss_only_mode}
            }

        # 3단계: 영상 상세 정보 조회
        # RSS 모드에서도 롱폼/쇼츠 구분을 위해 영상 길이 정보는 조회함
        print("3단계: 영상 정보 조회 중...")
        eel.update_progress("영상 정보 조회 중...", 75)()
        video_ids = [v['videoId'] for v in all_videos]
        video_info = get_videos_batch(youtube_service, video_ids)

        # 취소 확인
        if search_cancelled:
            return {'success': False, 'error': '검색이 중단되었습니다.', 'cancelled': True}

        # 4단계: 필터링
        print("4단계: 필터 적용 중...")
        eel.update_progress("필터 적용 중...", 90)()

        filtered_videos = []

        # 영상 타입에 따른 길이 필터 설정
        # 쇼츠: 183초(3분 3초) 이하, 롱폼: 184초 이상
        if video_type == 'shorts':
            min_duration = 0
            max_duration = 183
        else:  # 'long'
            min_duration = 184
            max_duration = float('inf')

        for video in all_videos:
            video_id = video['videoId']
            channel_id = video['channelId']

            v_info = video_info.get(video_id)
            if not v_info:
                continue

            # 영상 길이 필터 (롱폼/쇼츠 구분)
            duration = v_info.get('duration')
            if duration is None:
                duration = 0
            if duration < min_duration or duration > max_duration:
                continue

            # view_count가 None이면 0으로 처리
            view_count = v_info.get('viewCount')
            if view_count is None:
                view_count = 0

            # likeCount가 None이면 0으로 처리
            like_count = v_info.get('likeCount')
            if like_count is None:
                like_count = 0

            c_info = channel_info.get(channel_id)
            if not c_info:
                continue

            # subscriber_count가 None이면 0으로 처리
            subscriber_count = c_info.get('subscriberCount')
            if subscriber_count is None:
                subscriber_count = 0

            # 필터 타입별 조건 적용 (RSS 모드에서도 동일하게 적용)
            if filter_type == 'channel-monitor':
                # 채널모니터: 구독자 수 이하 & 조회수 이상
                if max_subscribers is not None and max_subscribers != float('inf'):
                    if subscriber_count > max_subscribers:
                        continue
                if view_count < min_views:
                    continue

            elif filter_type == 'keyword-search':
                # 키워드검색: 제목에 키워드 포함 & 조회수 이상
                if keyword and keyword.lower() not in video['title'].lower():
                    continue
                if view_count < min_views:
                    continue

            elif filter_type == 'hot-trend':
                # 핫트렌드: 최소 조회수 이상
                if view_count < min_views:
                    continue

            elif filter_type == 'mutation':
                # 돌연변이: 구독자 대비 조회수 비율
                if subscriber_count == 0:
                    continue
                ratio = view_count / subscriber_count
                if ratio < mutation_ratio:
                    continue

            filtered_videos.append({
                'videoId': video_id,
                'title': video['title'],
                'channelId': channel_id,
                'channelTitle': c_info.get('title', video.get('channelTitle', '알 수 없음')),
                'thumbnail': video['thumbnail'],
                'publishedAt': video['publishedAt'],
                'viewCount': view_count,
                'likeCount': like_count,
                'subscriberCount': subscriber_count,
                'duration': duration,
                'ratio': round(view_count / subscriber_count, 2) if subscriber_count > 0 else 0
            })

        # 정렬: RSS 모드는 날짜순, 핫트렌드는 조회수순, 그 외는 돌연변이 지수순 or 조회수순
        if rss_only_mode:
            # RSS 모드는 최신순 정렬
            filtered_videos.sort(key=lambda x: x['publishedAt'], reverse=True)
        elif filter_type == 'hot-trend':
            filtered_videos.sort(key=lambda x: x['viewCount'], reverse=True)
        elif filter_type == 'mutation':
            filtered_videos.sort(key=lambda x: x['ratio'], reverse=True)
        else:
            filtered_videos.sort(key=lambda x: x['viewCount'], reverse=True)

        eel.update_progress("완료!", 100)()
        print(f"필터링 결과: {len(filtered_videos)}개 (RSS 모드: {rss_only_mode})")

        return {
            'success': True,
            'videos': filtered_videos,
            'stats': {
                'total': len(all_videos),
                'filtered': len(filtered_videos),
                'rssMode': rss_only_mode
            }
        }

    except Exception as e:
        print(f"검색 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def clear_cache():
    """모든 캐시를 삭제합니다."""
    cache_manager.clear_all_cache()
    return {'success': True}


@eel.expose
def get_subscriptions_list():
    """구독 채널 목록을 반환합니다 (팝업용)."""
    global subscriptions
    return subscriptions


@eel.expose
def get_data_folder():
    """데이터 저장 폴더 경로를 반환합니다."""
    return DATA_DIR


@eel.expose
def export_subscriptions():
    """구독 목록을 JSON 파일로 내보냅니다 (파일 선택 다이얼로그)."""
    global subscriptions

    if not subscriptions:
        return {'success': False, 'error': '내보낼 구독 목록이 없습니다.'}

    try:
        # 파일 저장 다이얼로그
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        file_path = filedialog.asksaveasfilename(
            title='구독 목록 내보내기',
            defaultextension='.json',
            filetypes=[('JSON 파일', '*.json')],
            initialfile='subscriptions_export.json'
        )
        root.destroy()

        if not file_path:
            return {'success': False, 'error': '취소됨'}

        export_data = {
            'exportedAt': import_datetime().isoformat(),
            'count': len(subscriptions),
            'channels': [
                {
                    'id': sub['id'],
                    'title': sub['title'],
                    'thumbnail': sub.get('thumbnail', '')
                }
                for sub in subscriptions
            ]
        }

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)

        print(f"구독 목록 내보내기 완료: {file_path}")
        return {'success': True, 'path': file_path, 'count': len(subscriptions)}

    except Exception as e:
        print(f"내보내기 오류: {e}")
        return {'success': False, 'error': str(e)}


def import_datetime():
    """datetime 모듈에서 현재 시간 반환"""
    from datetime import datetime
    return datetime.now()


@eel.expose
def import_subscriptions():
    """JSON 파일에서 구독 목록을 가져와 일괄 구독합니다 (파일 선택 다이얼로그)."""
    global youtube_service

    try:
        # 파일 열기 다이얼로그
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        file_path = filedialog.askopenfilename(
            title='구독 목록 가져오기',
            filetypes=[('JSON 파일', '*.json')],
            initialfile='subscriptions_export.json'
        )
        root.destroy()

        if not file_path:
            return {'success': False, 'error': '취소됨'}

        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service:
            return {'success': False, 'error': '로그인이 필요합니다.'}

        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        channels = data.get('channels', [])
        if not channels:
            return {'success': False, 'error': '가져올 채널이 없습니다.'}

        # 현재 구독 목록 조회 (중복 방지)
        current_subs = set()
        try:
            request = youtube_service.subscriptions().list(
                part='snippet',
                mine=True,
                maxResults=50
            )
            while request:
                response = request.execute()
                for item in response.get('items', []):
                    current_subs.add(item['snippet']['resourceId']['channelId'])
                request = youtube_service.subscriptions().list_next(request, response)
        except Exception as e:
            print(f"현재 구독 목록 조회 실패: {e}")

        success_count = 0
        skip_count = 0
        fail_count = 0
        total = len(channels)

        for i, channel in enumerate(channels):
            channel_id = channel['id']
            channel_title = channel.get('title', channel_id)

            # 진행률 업데이트
            percent = int((i / total) * 100)
            eel.update_progress(f"구독 중: {channel_title} ({i+1}/{total})", percent)()

            # 이미 구독 중이면 건너뛰기
            if channel_id in current_subs:
                print(f"이미 구독 중: {channel_title}")
                skip_count += 1
                continue

            try:
                # 구독 추가
                youtube_service.subscriptions().insert(
                    part='snippet',
                    body={
                        'snippet': {
                            'resourceId': {
                                'kind': 'youtube#channel',
                                'channelId': channel_id
                            }
                        }
                    }
                ).execute()

                success_count += 1
                current_subs.add(channel_id)
                print(f"구독 완료: {channel_title}")

                # API 속도 제한 방지 (0.5초 딜레이)
                import time
                time.sleep(0.5)

            except Exception as e:
                fail_count += 1
                print(f"구독 실패 ({channel_title}): {e}")

        eel.update_progress("완료!", 100)()

        return {
            'success': True,
            'total': total,
            'subscribed': success_count,
            'skipped': skip_count,
            'failed': fail_count
        }

    except Exception as e:
        print(f"가져오기 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def unsubscribe_channel(channel_id):
    """채널 구독을 취소합니다."""
    global youtube_service, subscriptions

    try:
        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service:
            return {'success': False, 'error': '로그인이 필요합니다.'}

        # 구독 ID 찾기 (subscriptions.list로 조회)
        request = youtube_service.subscriptions().list(
            part='id',
            forChannelId=channel_id,
            mine=True
        )
        response = request.execute()

        if not response.get('items'):
            return {'success': False, 'error': '구독 정보를 찾을 수 없습니다.'}

        subscription_id = response['items'][0]['id']

        # 구독 취소
        youtube_service.subscriptions().delete(id=subscription_id).execute()

        # 로컬 목록에서도 제거
        subscriptions = [s for s in subscriptions if s['id'] != channel_id]

        # 캐시 업데이트
        cache_manager.save_subscriptions(subscriptions)

        print(f"채널 구독 취소 완료: {channel_id}")
        return {'success': True}

    except Exception as e:
        print(f"구독 취소 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def unsubscribe_channels_batch(channel_ids):
    """여러 채널의 구독을 일괄 취소합니다."""
    global youtube_service, subscriptions
    import time

    if not channel_ids:
        return {'success': False, 'error': '취소할 채널이 없습니다.'}

    try:
        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service:
            return {'success': False, 'error': '로그인이 필요합니다.'}

        success_count = 0
        fail_count = 0
        total = len(channel_ids)

        for i, channel_id in enumerate(channel_ids):
            # 진행률 업데이트
            percent = int((i / total) * 100)
            eel.update_progress(f"구독 취소 중: {i+1}/{total}", percent)()

            try:
                # 구독 ID 찾기
                request = youtube_service.subscriptions().list(
                    part='id',
                    forChannelId=channel_id,
                    mine=True
                )
                response = request.execute()

                if not response.get('items'):
                    fail_count += 1
                    continue

                subscription_id = response['items'][0]['id']

                # 구독 취소
                youtube_service.subscriptions().delete(id=subscription_id).execute()

                # 로컬 목록에서도 제거
                subscriptions = [s for s in subscriptions if s['id'] != channel_id]

                success_count += 1
                print(f"채널 구독 취소 완료: {channel_id}")

                # API 속도 제한 방지
                time.sleep(0.3)

            except Exception as e:
                fail_count += 1
                print(f"구독 취소 실패 ({channel_id}): {e}")

        # 캐시 업데이트
        cache_manager.save_subscriptions(subscriptions)

        eel.update_progress("완료!", 100)()

        return {
            'success': True,
            'total': total,
            'unsubscribed': success_count,
            'failed': fail_count
        }

    except Exception as e:
        print(f"일괄 구독 취소 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def select_excel_files():
    """엑셀 파일 선택 다이얼로그를 엽니다."""
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        file_paths = filedialog.askopenfilenames(
            title='엑셀 파일 선택',
            filetypes=[('Excel 파일', '*.xlsx *.xls')]
        )
        root.destroy()

        if not file_paths:
            return {'success': False, 'error': '취소됨'}

        return {'success': True, 'files': list(file_paths)}

    except Exception as e:
        print(f"파일 선택 오류: {e}")
        return {'success': False, 'error': str(e)}


def _parse_cell_ranges(cell_range_str):
    """
    셀 범위 문자열을 파싱합니다.
    여러 범위는 쉼표로 구분 (예: "A1:B5, A10:B15")

    Returns:
        list: [(start_col, start_row, end_col, end_row), ...]
    """
    import re
    ranges = []

    # 쉼표로 분리
    range_parts = [r.strip() for r in cell_range_str.split(',')]

    for part in range_parts:
        if not part:
            continue

        range_match = re.match(r'^([A-Z]+)(\d+):([A-Z]+)(\d+)$', part.upper().strip())
        if not range_match:
            continue

        start_col = range_match.group(1)
        start_row = int(range_match.group(2))
        end_col = range_match.group(3)
        end_row = int(range_match.group(4))

        ranges.append((start_col, start_row, end_col, end_row))

    return ranges


def _col_letter_to_num(col_letter):
    """열 문자를 숫자로 변환 (A=1, B=2, ..., Z=26, AA=27, ...)"""
    result = 0
    for char in col_letter:
        result = result * 26 + (ord(char) - ord('A') + 1)
    return result


@eel.expose
def extract_urls_from_excel(file_paths, cell_range):
    """
    엑셀 파일들에서 지정된 셀 범위의 URL을 추출합니다.
    여러 범위를 쉼표로 구분하여 입력 가능.

    Args:
        file_paths: 엑셀 파일 경로 리스트
        cell_range: 셀 범위 (예: "A2:A100" 또는 "A1:B5, A10:B15")

    Returns:
        dict: {'success': bool, 'urls': [...], 'file_count': int}
    """
    try:
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter

        # 셀 범위 파싱
        ranges = _parse_cell_ranges(cell_range)
        if not ranges:
            return {'success': False, 'error': f'잘못된 셀 범위 형식: {cell_range}\n예: A2:A100 또는 A1:B5, A10:B15'}

        all_urls = []
        file_results = []

        for file_path in file_paths:
            try:
                wb = load_workbook(file_path, read_only=True, data_only=True)
                ws = wb.active

                file_urls = []

                for start_col, start_row, end_col, end_row in ranges:
                    start_col_num = _col_letter_to_num(start_col)
                    end_col_num = _col_letter_to_num(end_col)

                    for row in range(start_row, end_row + 1):
                        for col_num in range(start_col_num, end_col_num + 1):
                            col_letter = get_column_letter(col_num)
                            cell_ref = f"{col_letter}{row}"
                            cell_value = ws[cell_ref].value

                            if cell_value and str(cell_value).strip():
                                url = str(cell_value).strip()
                                file_urls.append(url)
                                all_urls.append(url)

                wb.close()

                file_results.append({
                    'file': file_path.split('\\')[-1].split('/')[-1],
                    'count': len(file_urls)
                })

            except Exception as e:
                file_results.append({
                    'file': file_path.split('\\')[-1].split('/')[-1],
                    'error': str(e)
                })

        # 중복 제거
        unique_urls = list(dict.fromkeys(all_urls))

        return {
            'success': True,
            'urls': unique_urls,
            'total_count': len(all_urls),
            'unique_count': len(unique_urls),
            'file_count': len(file_paths),
            'file_results': file_results
        }

    except Exception as e:
        print(f"엑셀 URL 추출 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def extract_cells_from_excel(file_paths, cell_range):
    """
    엑셀 파일들에서 지정된 셀 범위의 데이터를 추출합니다.
    URL 추출과 별도로 순수 셀 데이터만 추출합니다.

    Args:
        file_paths: 엑셀 파일 경로 리스트
        cell_range: 셀 범위 (예: "A1:B5, A10:B15")

    Returns:
        dict: {'success': bool, 'data': [...], 'total_cells': int}
    """
    try:
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter

        # 셀 범위 파싱
        ranges = _parse_cell_ranges(cell_range)
        if not ranges:
            return {'success': False, 'error': f'잘못된 셀 범위 형식: {cell_range}\n예: A1:B5, A10:B15'}

        all_data = []
        file_results = []

        for file_path in file_paths:
            try:
                wb = load_workbook(file_path, read_only=True, data_only=True)
                ws = wb.active
                filename = file_path.split('\\')[-1].split('/')[-1]

                file_data = []

                for start_col, start_row, end_col, end_row in ranges:
                    start_col_num = _col_letter_to_num(start_col)
                    end_col_num = _col_letter_to_num(end_col)

                    range_data = []
                    for row in range(start_row, end_row + 1):
                        row_data = []
                        for col_num in range(start_col_num, end_col_num + 1):
                            col_letter = get_column_letter(col_num)
                            cell_ref = f"{col_letter}{row}"
                            cell_value = ws[cell_ref].value

                            if cell_value is not None:
                                row_data.append(str(cell_value))
                            else:
                                row_data.append('')

                        # 빈 행이 아니면 추가
                        if any(cell.strip() for cell in row_data):
                            range_data.append(row_data)

                    if range_data:
                        file_data.append({
                            'range': f'{start_col}{start_row}:{end_col}{end_row}',
                            'rows': range_data
                        })

                wb.close()

                total_cells = sum(len(rd['rows']) for rd in file_data)
                file_results.append({
                    'file': filename,
                    'ranges': file_data,
                    'cell_count': total_cells
                })

                all_data.extend(file_data)

            except Exception as e:
                file_results.append({
                    'file': file_path.split('\\')[-1].split('/')[-1],
                    'error': str(e)
                })

        total_cells = sum(fr.get('cell_count', 0) for fr in file_results if 'cell_count' in fr)

        return {
            'success': True,
            'file_results': file_results,
            'file_count': len(file_paths),
            'total_cells': total_cells
        }

    except Exception as e:
        print(f"엑셀 셀 추출 오류: {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def resolve_channel_urls(urls):
    """
    URL 목록에서 채널 ID를 조회합니다.

    Args:
        urls: URL/핸들 리스트

    Returns:
        dict: {'success': bool, 'channels': [...], 'failed': [...]}
    """
    global youtube_service

    try:
        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service:
            return {'success': False, 'error': '로그인이 필요합니다.'}

        from youtube_api import resolve_channel_ids_batch

        total = len(urls)
        resolved = {'success': [], 'failed': []}

        def progress_callback(current, total_count, url, result):
            percent = int((current / total_count) * 50)
            eel.update_progress(f"채널 조회: {current}/{total_count}", percent)()

        resolved = resolve_channel_ids_batch(youtube_service, urls, progress_callback)

        return {
            'success': True,
            'channels': resolved['success'],
            'failed': resolved['failed'],
            'success_count': len(resolved['success']),
            'failed_count': len(resolved['failed'])
        }

    except Exception as e:
        print(f"채널 조회 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def subscribe_channels_from_urls(channel_ids):
    """
    채널 ID 목록으로 일괄 구독합니다.
    이미 구독한 채널은 API 호출 없이 건너뛰어 할당량을 절약합니다.

    Args:
        channel_ids: 채널 ID 리스트

    Returns:
        dict: {'success': bool, 'subscribed': int, 'already': int, 'failed': int}
    """
    global youtube_service, subscriptions

    if not channel_ids:
        return {'success': False, 'error': '구독할 채널이 없습니다.'}

    try:
        if not youtube_service:
            youtube_service = get_authenticated_service()

        if not youtube_service:
            return {'success': False, 'error': '로그인이 필요합니다.'}

        import time

        # 현재 구독 중인 채널 ID 목록 가져오기 (할당량 절약을 위해)
        eel.update_progress("구독 목록 확인 중...", 0)()
        existing_channel_ids = set()

        # 이미 로드된 구독 목록이 있으면 사용
        if subscriptions:
            existing_channel_ids = {sub.get('channelId') for sub in subscriptions if sub.get('channelId')}
        else:
            # 없으면 API로 가져오기
            try:
                page_token = None
                while True:
                    response = youtube_service.subscriptions().list(
                        part='snippet',
                        mine=True,
                        maxResults=50,
                        pageToken=page_token
                    ).execute()

                    for item in response.get('items', []):
                        channel_id = item['snippet']['resourceId']['channelId']
                        existing_channel_ids.add(channel_id)

                    page_token = response.get('nextPageToken')
                    if not page_token:
                        break
            except Exception as e:
                print(f"구독 목록 조회 실패, 중복 체크 없이 진행: {e}")

        # 이미 구독한 채널 필터링
        new_channel_ids = [cid for cid in channel_ids if cid not in existing_channel_ids]
        already_subscribed = len(channel_ids) - len(new_channel_ids)

        print(f"총 {len(channel_ids)}개 중 {already_subscribed}개 이미 구독, {len(new_channel_ids)}개 새로 구독 예정")

        if not new_channel_ids:
            eel.update_progress("완료!", 100)()
            return {
                'success': True,
                'total': len(channel_ids),
                'subscribed': 0,
                'already': already_subscribed,
                'failed': 0,
                'message': '모든 채널이 이미 구독되어 있습니다.'
            }

        total = len(new_channel_ids)
        subscribed = 0
        failed = 0

        for i, channel_id in enumerate(new_channel_ids):
            percent = int(((i + 1) / total) * 100)
            eel.update_progress(f"구독 중: {i+1}/{total} (건너뜀: {already_subscribed})", percent)()

            try:
                youtube_service.subscriptions().insert(
                    part='snippet',
                    body={
                        'snippet': {
                            'resourceId': {
                                'kind': 'youtube#channel',
                                'channelId': channel_id
                            }
                        }
                    }
                ).execute()
                subscribed += 1

            except Exception as e:
                error_msg = str(e)
                if 'subscriptionDuplicate' in error_msg:
                    # 혹시 중복이면 already에 추가
                    already_subscribed += 1
                else:
                    failed += 1
                    print(f"구독 실패 ({channel_id}): {e}")

            # API 속도 제한 방지
            time.sleep(0.3)

        eel.update_progress("완료!", 100)()

        return {
            'success': True,
            'total': len(channel_ids),
            'subscribed': subscribed,
            'already': already_subscribed,
            'failed': failed
        }

    except Exception as e:
        print(f"일괄 구독 오류: {e}")
        return {'success': False, 'error': str(e)}


# ===================== 멀티 계정 관리 함수 =====================

@eel.expose
def get_accounts():
    """
    저장된 모든 계정 목록을 반환합니다.

    Returns:
        dict: {'success': bool, 'accounts': [...], 'current_account_id': str}
    """
    try:
        accounts_data = account_manager.load_accounts()

        # 각 계정의 인증 상태 및 API 설정 상태 확인
        for account in accounts_data['accounts']:
            has_api = account_manager.has_account_api_credentials(account['id'])
            account['has_api_config'] = has_api
            # API 설정이 있어야 인증 상태 확인 가능
            if has_api:
                account['is_authenticated'] = is_authenticated(account['id'])
            else:
                account['is_authenticated'] = False

        return {
            'success': True,
            'accounts': accounts_data['accounts'],
            'current_account_id': accounts_data['current_account_id']
        }
    except Exception as e:
        print(f"계정 목록 조회 오류: {e}")
        return {'success': False, 'accounts': [], 'current_account_id': None, 'error': str(e)}


@eel.expose
def create_first_account(api_key, client_id, client_secret):
    """
    첫 번째 계정을 생성합니다 (API 설정 포함).

    Args:
        api_key: YouTube Data API v3 키
        client_id: Google OAuth Client ID
        client_secret: Google OAuth Client Secret

    Returns:
        dict: {'success': bool, 'account_id': str, 'error': str}
    """
    try:
        # 새 계정 생성
        result = account_manager.add_account(
            name='새 계정',
            email='',
            thumbnail=''
        )

        if not result['success']:
            return result

        account_id = result['account_id']

        # API 자격 증명 저장
        api_result = account_manager.save_account_api_credentials(
            account_id, api_key.strip(), client_id.strip(), client_secret.strip()
        )

        if not api_result['success']:
            # 실패 시 계정 삭제
            account_manager.remove_account(account_id)
            return {'success': False, 'error': f'API 설정 저장 실패: {api_result.get("error", "")}'}

        # 런타임 config에 적용
        config.set_current_credentials(api_key.strip(), client_id.strip(), client_secret.strip())

        print(f"첫 번째 계정 생성 완료: {account_id}")
        return {
            'success': True,
            'account_id': account_id,
            'needsLogin': True
        }

    except Exception as e:
        print(f"첫 계정 생성 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def add_new_account():
    """
    새 계정 추가를 준비합니다.
    실제 계정은 API 설정과 함께 add_account_with_api로 추가됩니다.

    Returns:
        dict: {'success': bool, 'needsApiSetup': bool}
    """
    # 계정 추가 전에 API 설정이 필요함을 알림
    return {
        'success': True,
        'needsApiSetup': True
    }


@eel.expose
def add_account_with_api(api_key, client_id, client_secret):
    """
    새 계정을 API 설정과 함께 추가합니다.

    Args:
        api_key: YouTube Data API v3 키
        client_id: Google OAuth Client ID
        client_secret: Google OAuth Client Secret

    Returns:
        dict: {'success': bool, 'account_id': str, 'error': str}
    """
    try:
        # 새 계정 생성
        result = account_manager.add_account(
            name='새 계정',
            email='',
            thumbnail=''
        )

        if not result['success']:
            return result

        account_id = result['account_id']

        # API 자격 증명 저장
        api_result = account_manager.save_account_api_credentials(
            account_id, api_key.strip(), client_id.strip(), client_secret.strip()
        )

        if not api_result['success']:
            # 실패 시 계정 삭제
            account_manager.remove_account(account_id)
            return {'success': False, 'error': f'API 설정 저장 실패: {api_result.get("error", "")}'}

        # 런타임 config에 적용 (새 계정으로 로그인할 예정이므로)
        config.set_current_credentials(api_key.strip(), client_id.strip(), client_secret.strip())

        print(f"계정 추가 완료 (API 설정 포함): {account_id}")
        return {
            'success': True,
            'account_id': account_id,
            'needsLogin': True
        }

    except Exception as e:
        print(f"계정 추가 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def login_account(account_id):
    """
    특정 계정으로 로그인합니다 (브라우저 OAuth).

    Args:
        account_id: 로그인할 계정 ID

    Returns:
        dict: {'success': bool, 'error': str}
    """
    global youtube_service

    if not is_configured():
        return {'success': False, 'error': 'API 설정이 필요합니다.'}

    try:
        # 사용 가능한 포트 찾기
        port = find_auth_port(8888)
        print(f"인증 서버 포트: {port}")

        # 인증 URL 생성
        flow, auth_url = get_auth_url_with_localhost(port)
        if not flow or not auth_url:
            return {'success': False, 'error': 'API 설정을 확인해주세요.'}

        # 전용 Chrome 창에서 인증 URL 열기
        print("전용 Chrome 창에서 로그인 페이지를 엽니다...")
        open_auth_browser(auth_url)

        # 인증 콜백 서버 시작 (블로킹)
        print("인증 대기 중...")
        auth_result = start_auth_server(port)

        if auth_result.get('error'):
            return {'success': False, 'error': f"로그인 실패: {auth_result['error']}"}

        if not auth_result.get('code'):
            return {'success': False, 'error': '인증 코드를 받지 못했습니다.'}

        # 코드로 토큰 교환 (특정 계정 ID로 저장)
        print("토큰 교환 중...")
        creds = exchange_code_for_token(flow, auth_result['code'], account_id)

        if not creds:
            return {'success': False, 'error': '토큰 교환에 실패했습니다.'}

        # YouTube 서비스 생성
        from googleapiclient.discovery import build
        youtube_service = build('youtube', 'v3', credentials=creds)

        # 채널 정보 조회하여 계정 정보 업데이트
        try:
            request = youtube_service.channels().list(
                part='snippet',
                mine=True
            )
            response = request.execute()

            if response.get('items'):
                item = response['items'][0]
                account_manager.update_account(
                    account_id,
                    name=item['snippet']['title'],
                    email=item['id'],
                    thumbnail=item['snippet']['thumbnails']['default']['url']
                )
        except Exception as e:
            print(f"채널 정보 조회 실패: {e}")

        # 현재 계정으로 전환
        account_manager.switch_account(account_id)

        print(f"계정 {account_id} 로그인 성공!")
        return {'success': True}

    except Exception as e:
        print(f"계정 로그인 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def switch_to_account(account_id):
    """
    다른 계정으로 전환합니다.

    각 계정은 반드시 자체 API 설정이 있어야 합니다.

    Args:
        account_id: 전환할 계정 ID

    Returns:
        dict: {'success': bool, 'account': dict, 'error': str}
    """
    global youtube_service, subscriptions

    try:
        # 계정별 API 자격 증명이 있는지 확인
        has_own_api = account_manager.has_account_api_credentials(account_id)

        if not has_own_api:
            # API 설정이 없으면 오류
            return {
                'success': False,
                'error': '이 계정에 API 설정이 없습니다. API 설정을 먼저 해주세요.',
                'needsApiSetup': True
            }

        # API 자격 증명 로드
        api_result = account_manager.load_account_api_credentials(account_id)
        if not api_result['success']:
            return {'success': False, 'error': api_result.get('error', 'API 설정 로드 실패')}

        # 런타임 config에 적용
        config.set_current_credentials(api_result['api_key'], api_result['client_id'], api_result['client_secret'])
        print(f"계정 {account_id}의 API 자격 증명 로드 완료")

        # 인증 상태 확인
        if not is_authenticated(account_id):
            return {'success': False, 'error': '해당 계정은 재로그인이 필요합니다.', 'needsLogin': True}

        # 계정 전환
        result = account_manager.switch_account(account_id)
        if not result['success']:
            return result

        # YouTube 서비스 재생성
        youtube_service = get_authenticated_service(account_id)
        if not youtube_service:
            return {'success': False, 'error': '서비스 연결에 실패했습니다.'}

        # 구독 목록 및 캐시 초기화
        subscriptions = []
        cache_manager.clear_all_cache()

        print(f"계정 전환 완료: {result['account']['name']}")
        return {'success': True, 'account': result['account']}

    except Exception as e:
        print(f"계정 전환 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def remove_account_by_id(account_id):
    """
    계정을 삭제합니다.

    Args:
        account_id: 삭제할 계정 ID

    Returns:
        dict: {'success': bool, 'error': str}
    """
    global youtube_service, subscriptions

    try:
        # 현재 계정인지 확인
        current = account_manager.get_current_account()
        is_current = current and current['id'] == account_id

        # 계정 삭제
        result = account_manager.remove_account(account_id)

        if result['success'] and is_current:
            # 현재 계정이 삭제된 경우 서비스 재생성
            new_current = account_manager.get_current_account()
            if new_current:
                youtube_service = get_authenticated_service(new_current['id'])
            else:
                youtube_service = None

            # 구독 목록 및 캐시 초기화
            subscriptions = []
            cache_manager.clear_all_cache()

        return result

    except Exception as e:
        print(f"계정 삭제 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def rename_account(account_id, new_name):
    """
    계정 이름을 변경합니다.

    Args:
        account_id: 계정 ID
        new_name: 새 이름

    Returns:
        dict: {'success': bool, 'error': str}
    """
    try:
        return account_manager.update_account(account_id, name=new_name)
    except Exception as e:
        print(f"계정 이름 변경 오류: {e}")
        return {'success': False, 'error': str(e)}


# ===================== 계정별 API 자격 증명 관리 =====================

@eel.expose
def save_account_api_config(account_id, api_key, client_id, client_secret):
    """
    계정별 API 자격 증명을 저장합니다.

    Args:
        account_id: 계정 ID
        api_key: YouTube Data API v3 키
        client_id: Google OAuth Client ID
        client_secret: Google OAuth Client Secret

    Returns:
        dict: {'success': bool, 'error': str}
    """
    try:
        result = account_manager.save_account_api_credentials(
            account_id, api_key.strip(), client_id.strip(), client_secret.strip()
        )

        if result['success']:
            # 현재 계정이면 런타임 config에도 적용
            current = account_manager.get_current_account()
            if current and current['id'] == account_id:
                config.set_current_credentials(api_key.strip(), client_id.strip(), client_secret.strip())

        return result
    except Exception as e:
        print(f"계정 API 설정 저장 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def save_account_api_key(account_id, api_key):
    """
    기존 계정에 API 키만 저장/업데이트합니다.
    (Client ID와 Client Secret은 기존 값 유지)

    Args:
        account_id: 계정 ID
        api_key: YouTube Data API v3 키

    Returns:
        dict: {'success': bool, 'error': str}
    """
    try:
        # 기존 API 자격 증명 로드
        existing = account_manager.load_account_api_credentials(account_id)

        if existing['success']:
            # 기존 Client ID/Secret 유지하고 API 키만 업데이트
            client_id = existing['client_id']
            client_secret = existing['client_secret']
        else:
            # 기존 설정이 없으면 현재 런타임의 값 사용
            client_id = config.get_client_id()
            client_secret = config.get_client_secret()

            if not client_id or not client_secret:
                return {'success': False, 'error': 'Client ID와 Client Secret이 없습니다. 먼저 로그인해주세요.'}

        result = account_manager.save_account_api_credentials(
            account_id, api_key.strip(), client_id, client_secret
        )

        if result['success']:
            # 현재 계정이면 런타임 config에도 적용
            current = account_manager.get_current_account()
            if current and current['id'] == account_id:
                config.set_current_credentials(api_key.strip(), client_id, client_secret)

        return result
    except Exception as e:
        print(f"계정 API 키 저장 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def load_account_api_config(account_id):
    """
    계정별 API 자격 증명을 불러옵니다.

    Args:
        account_id: 계정 ID

    Returns:
        dict: {'success': bool, 'api_key': str, 'client_id': str, 'client_secret': str, 'error': str}
    """
    try:
        result = account_manager.load_account_api_credentials(account_id)

        if result['success']:
            # 현재 계정이면 런타임 config에 적용
            current = account_manager.get_current_account()
            if current and current['id'] == account_id:
                config.set_current_credentials(result['api_key'], result['client_id'], result['client_secret'])

        return result
    except Exception as e:
        print(f"계정 API 설정 로드 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def has_account_api_config(account_id):
    """
    계정에 API 자격 증명이 저장되어 있는지 확인합니다.

    Args:
        account_id: 계정 ID

    Returns:
        dict: {'has_config': bool}
    """
    try:
        has_config = account_manager.has_account_api_credentials(account_id)
        return {'has_config': has_config}
    except Exception as e:
        print(f"계정 API 설정 확인 오류: {e}")
        return {'has_config': False}


@eel.expose
def delete_account_api_config(account_id):
    """
    계정의 API 자격 증명을 삭제합니다.

    Args:
        account_id: 계정 ID

    Returns:
        dict: {'success': bool}
    """
    try:
        success = account_manager.delete_account_api_credentials(account_id)

        if success:
            # 현재 계정이면 런타임 config 초기화
            current = account_manager.get_current_account()
            if current and current['id'] == account_id:
                config.clear_current_credentials()

        return {'success': success}
    except Exception as e:
        print(f"계정 API 설정 삭제 오류: {e}")
        return {'success': False}


@eel.expose
def verify_account_api_password(account_id, password):
    """
    계정 API 비밀번호가 맞는지 확인합니다.

    Args:
        account_id: 계정 ID
        password: 확인할 비밀번호

    Returns:
        dict: {'valid': bool}
    """
    try:
        valid = account_manager.verify_account_api_password(account_id, password)
        return {'valid': valid}
    except Exception as e:
        print(f"비밀번호 확인 오류: {e}")
        return {'valid': False}


@eel.expose
def get_account_api_status(account_id):
    """
    계정의 API 설정 상태를 반환합니다.

    Args:
        account_id: 계정 ID

    Returns:
        dict: {'has_own_api': bool, 'needs_setup': bool}
    """
    try:
        has_own = account_manager.has_account_api_credentials(account_id)
        return {
            'has_own_api': has_own,
            'needs_setup': not has_own  # API 설정이 없으면 설정 필요
        }
    except Exception as e:
        print(f"API 상태 확인 오류: {e}")
        return {'has_own_api': False, 'needs_setup': True}


@eel.expose
def get_current_account_info():
    """
    현재 선택된 계정 정보를 반환합니다.

    Returns:
        dict: {'success': bool, 'account': dict}
    """
    try:
        account = account_manager.get_current_account()
        if account:
            account['is_authenticated'] = is_authenticated(account['id'])
            return {'success': True, 'account': account}
        return {'success': False, 'account': None}
    except Exception as e:
        print(f"현재 계정 조회 오류: {e}")
        return {'success': False, 'account': None, 'error': str(e)}


@eel.expose
def get_video_comments_filtered(video_id, keywords=None, max_count=20):
    """
    영상의 댓글을 가져옵니다 (키워드 필터링).

    Args:
        video_id: 영상 ID
        keywords: 필터 키워드 리스트
        max_count: 최대 댓글 수

    Returns:
        dict: {'success': bool, 'comments': list}
    """
    global youtube_service

    print(f"[댓글 조회] video_id={video_id}, keywords={keywords}, max_count={max_count}")

    if not youtube_service:
        print("[댓글 조회] 로그인 필요")
        return {'success': False, 'error': '로그인이 필요합니다', 'comments': []}

    try:
        comments = get_filtered_comments(youtube_service, video_id, keywords, max_count)
        print(f"[댓글 조회] 성공: {len(comments)}개 댓글")
        return {'success': True, 'comments': comments}
    except Exception as e:
        print(f"[댓글 조회] 오류 ({video_id}): {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e), 'comments': []}


def on_close(page, sockets):
    """브라우저 창이 닫히면 프로그램 종료 (토큰은 유지 - 토큰생성기로 미리 만들어둔 토큰 보존)"""
    print("\n[종료] 브라우저 창이 닫혔습니다. 프로그램을 종료합니다...")
    # 토큰을 삭제하지 않고 유지 (토큰생성기로 미리 만들어둔 토큰을 보존)
    import sys
    sys.exit(0)


def find_free_port(start_port=8000):
    """사용 가능한 포트를 찾습니다."""
    import socket
    for port in range(start_port, start_port + 100):
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.bind(('localhost', port))
                return port
        except OSError:
            continue
    return start_port


# ===== MP3 추출 및 Whisper 텍스트 변환 기능 =====

# 지원 확장자
VIDEO_EXT = {'.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm', '.m4v', '.mpeg', '.mpg', '.3gp'}
AUDIO_EXT = {'.mp3', '.wav', '.m4a', '.flac', '.ogg', '.aac', '.wma'}

# Whisper 모델
whisper_model = None


def get_ffmpeg_path():
    """FFmpeg 경로를 찾습니다."""
    if sys.platform == 'win32':
        name = 'ffmpeg.exe'
    else:
        name = 'ffmpeg'

    # PATH에서 찾기
    for path in os.environ.get('PATH', '').split(os.pathsep):
        full = os.path.join(path, name)
        if os.path.isfile(full):
            return full

    # 실행 파일과 같은 폴더에서 찾기
    if getattr(sys, 'frozen', False):
        exe_dir = os.path.dirname(sys.executable)
    else:
        exe_dir = os.path.dirname(os.path.abspath(__file__))

    local = os.path.join(exe_dir, name)
    if os.path.isfile(local):
        return local

    return name


@eel.expose
def check_ffmpeg():
    """FFmpeg 설치 여부를 확인합니다."""
    try:
        result = subprocess.run(
            [get_ffmpeg_path(), '-version'],
            capture_output=True,
            creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
        )
        return result.returncode == 0
    except:
        return False


@eel.expose
def check_whisper():
    """Whisper 설치 여부를 확인합니다."""
    try:
        import whisper
        return True
    except:
        return False


@eel.expose
def load_whisper_model(model_name='small'):
    """Whisper 모델을 로드합니다."""
    global whisper_model
    try:
        import whisper
        whisper_model = whisper.load_model(model_name)
        return True
    except Exception as e:
        print(f"모델 로드 실패: {e}")
        return False


@eel.expose
def select_media_files():
    """미디어 파일 선택 다이얼로그를 엽니다."""
    from tkinter import Tk, filedialog
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)

    files = filedialog.askopenfilenames(
        title='파일 선택',
        filetypes=[
            ('미디어 파일', '*.mp4 *.avi *.mkv *.mov *.mp3 *.wav *.m4a *.flac'),
            ('동영상', '*.mp4 *.avi *.mkv *.mov *.wmv *.flv *.webm'),
            ('오디오', '*.mp3 *.wav *.m4a *.flac *.ogg *.aac'),
            ('모든 파일', '*.*')
        ]
    )
    root.destroy()
    return list(files)


@eel.expose
def select_media_folder():
    """미디어 폴더 선택 다이얼로그를 엽니다."""
    from tkinter import Tk, filedialog
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    folder = filedialog.askdirectory(title='폴더 선택')
    root.destroy()
    return folder


@eel.expose
def select_output_folder():
    """출력 폴더 선택 다이얼로그를 엽니다."""
    from tkinter import Tk, filedialog
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    folder = filedialog.askdirectory(title='출력 폴더 선택')
    root.destroy()
    return folder


@eel.expose
def get_media_files_from_folder(folder_path):
    """폴더에서 미디어 파일 목록을 가져옵니다."""
    if not folder_path or not os.path.isdir(folder_path):
        return []

    files = []
    for f in os.listdir(folder_path):
        ext = os.path.splitext(f)[1].lower()
        if ext in VIDEO_EXT or ext in AUDIO_EXT:
            files.append(os.path.join(folder_path, f))
    return sorted(files)


def is_video_file(path):
    """동영상 파일인지 확인합니다."""
    return os.path.splitext(path)[1].lower() in VIDEO_EXT


def extract_audio_from_video(video_path, output_path, bitrate='192', progress_callback=None):
    """동영상에서 오디오를 추출합니다."""
    try:
        import re

        cmd = [
            get_ffmpeg_path(),
            '-i', video_path,
            '-vn',
            '-acodec', 'libmp3lame',
            '-ab', f'{bitrate}k',
            '-ar', '44100',
            '-y',
            output_path
        ]

        process = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            universal_newlines=True,
            creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
        )

        # 영상 길이 추출
        duration = None
        duration_pattern = re.compile(r'Duration: (\d{2}):(\d{2}):(\d{2})\.\d+')
        time_pattern = re.compile(r'time=(\d{2}):(\d{2}):(\d{2})\.\d+')

        stderr_output = []

        for line in process.stderr:
            stderr_output.append(line)

            # Duration 추출
            if duration is None:
                match = duration_pattern.search(line)
                if match:
                    h, m, s = map(int, match.groups())
                    duration = h * 3600 + m * 60 + s

            # 현재 진행 시간 추출 및 진행률 계산
            if duration and progress_callback:
                match = time_pattern.search(line)
                if match:
                    h, m, s = map(int, match.groups())
                    current_time = h * 3600 + m * 60 + s
                    progress_percent = min(int((current_time / duration) * 100), 99)
                    try:
                        progress_callback(progress_percent)
                    except:
                        pass

        process.wait()

        # 완료 시 100% 전송
        if progress_callback:
            try:
                progress_callback(100)
            except:
                pass

        return process.returncode == 0, ''.join(stderr_output)
    except Exception as e:
        return False, str(e)


def transcribe_audio_file(audio_path, language='ko'):
    """오디오 파일을 텍스트로 변환합니다."""
    global whisper_model

    if whisper_model is None:
        return None, "모델이 로드되지 않았습니다"

    try:
        result = whisper_model.transcribe(audio_path, language=language if language != 'auto' else None)
        return result, None
    except Exception as e:
        return None, str(e)


def format_timestamp(seconds):
    """초를 타임스탬프 형식으로 변환합니다."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def save_transcription_result(result, output_path, fmt='txt'):
    """변환 결과를 파일로 저장합니다."""
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            if fmt == 'txt':
                f.write(result['text'])

            elif fmt == 'txt_timestamp':
                for seg in result['segments']:
                    start = format_timestamp(seg['start'])
                    end = format_timestamp(seg['end'])
                    f.write(f"[{start} --> {end}]\n{seg['text'].strip()}\n\n")

            elif fmt == 'srt':
                for i, seg in enumerate(result['segments'], 1):
                    start = format_timestamp(seg['start'])
                    end = format_timestamp(seg['end'])
                    f.write(f"{i}\n{start} --> {end}\n{seg['text'].strip()}\n\n")

            elif fmt == 'json':
                json.dump(result, f, ensure_ascii=False, indent=2)

        return True
    except:
        return False


@eel.expose
def start_media_processing(file_list, output_folder, options):
    """미디어 파일 처리를 시작합니다."""
    def process():
        total = len(file_list)
        extract_mp3 = options.get('extract_mp3', True)
        transcribe = options.get('transcribe', True)
        bitrate = options.get('bitrate', '192')
        language = options.get('language', 'ko')
        out_fmt = options.get('output_format', 'txt')
        same_folder = options.get('same_folder', True)

        for i, filepath in enumerate(file_list, 1):
            filename = os.path.basename(filepath)
            name_only = os.path.splitext(filename)[0]

            out_dir = os.path.dirname(filepath) if same_folder else output_folder

            eel.update_media_progress(i, total, filename, 'processing', '처리 중...')()

            audio_path = filepath

            # MP3 추출
            if is_video_file(filepath) and extract_mp3:
                mp3_path = os.path.join(out_dir, f'{name_only}.mp3')

                counter = 1
                while os.path.exists(mp3_path):
                    mp3_path = os.path.join(out_dir, f'{name_only}_{counter}.mp3')
                    counter += 1

                eel.update_media_progress(i, total, filename, 'processing', 'MP3 추출 중... 0%')()

                # 진행률 콜백 함수
                def on_extract_progress(percent):
                    try:
                        eel.update_media_progress(i, total, filename, 'processing', f'MP3 추출 중... {percent}%')()
                    except:
                        pass

                success, error = extract_audio_from_video(filepath, mp3_path, bitrate, on_extract_progress)
                if not success:
                    eel.update_media_progress(i, total, filename, 'error', f'MP3 실패')()
                    continue

                audio_path = mp3_path

            # 텍스트 변환
            if transcribe:
                eel.update_media_progress(i, total, filename, 'processing', '텍스트 변환 중...')()

                result, error = transcribe_audio_file(audio_path, language)
                if result is None:
                    print(f"[Whisper 에러] {filename}: {error}")
                    eel.update_media_progress(i, total, filename, 'error', f'변환 실패: {error}')()
                    continue

                ext_map = {'txt': '.txt', 'txt_timestamp': '_timestamp.txt', 'srt': '.srt', 'json': '.json'}
                txt_path = os.path.join(out_dir, f'{name_only}{ext_map.get(out_fmt, ".txt")}')

                counter = 1
                base = txt_path
                while os.path.exists(txt_path):
                    name_part = os.path.splitext(base)[0]
                    ext_part = os.path.splitext(base)[1]
                    txt_path = f'{name_part}_{counter}{ext_part}'
                    counter += 1

                save_transcription_result(result, txt_path, out_fmt)

            eel.update_media_progress(i, total, filename, 'done', '완료!')()

        eel.media_processing_complete()()

    thread = threading.Thread(target=process)
    thread.daemon = True
    thread.start()
    return {'success': True}


# ===== Word 파일 (.docx) 읽기 API =====

@eel.expose
def read_docx_file(file_path):
    """
    Word 파일(.docx)에서 텍스트를 추출합니다.

    Args:
        file_path: Word 파일 경로

    Returns:
        dict: {'success': bool, 'text': str, 'error': str}
    """
    try:
        from docx import Document

        if not os.path.exists(file_path):
            return {'success': False, 'error': '파일을 찾을 수 없습니다.'}

        if not file_path.lower().endswith('.docx'):
            return {'success': False, 'error': '.docx 파일만 지원됩니다.'}

        doc = Document(file_path)

        # 모든 문단의 텍스트 추출
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # 빈 문단 제외
                paragraphs.append(text)

        full_text = '\n'.join(paragraphs)

        return {
            'success': True,
            'text': full_text,
            'paragraphs': len(paragraphs)
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def select_and_read_docx_file():
    """
    파일 선택 대화상자를 열고 선택한 Word 파일의 텍스트를 반환합니다.

    Returns:
        dict: {'success': bool, 'text': str, 'filename': str, 'error': str}
    """
    try:
        import tkinter as tk
        from tkinter import filedialog
        from docx import Document

        # 파일 선택 대화상자
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        file_path = filedialog.askopenfilename(
            title='Word 파일 선택',
            filetypes=[
                ('Word 문서', '*.docx'),
                ('텍스트 파일', '*.txt'),
                ('모든 파일', '*.*')
            ]
        )

        root.destroy()

        if not file_path:
            return {'success': False, 'error': '파일 선택이 취소되었습니다.', 'cancelled': True}

        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)

        # 파일 확장자에 따라 처리
        if file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text
                if text.strip():
                    paragraphs.append(text)
            full_text = '\n'.join(paragraphs)
        elif file_path.lower().endswith('.txt'):
            # 텍스트 파일 읽기 (인코딩 자동 감지)
            encodings = ['utf-8', 'cp949', 'euc-kr', 'utf-16']
            full_text = None
            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        full_text = f.read()
                    break
                except:
                    continue
            if full_text is None:
                return {'success': False, 'error': '파일 인코딩을 인식할 수 없습니다.'}
        else:
            return {'success': False, 'error': '지원하지 않는 파일 형식입니다. (.docx 또는 .txt)'}

        return {
            'success': True,
            'text': full_text,
            'filename': filename,
            'fileSize': file_size,
            'fileType': 'docx' if file_path.lower().endswith('.docx') else 'txt',
            'path': file_path  # 자막 분할 기능용 파일 경로
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def select_and_read_multiple_files():
    """
    여러 파일을 선택하고 텍스트를 반환합니다 (텍스트합치기용).

    Returns:
        dict: {'success': bool, 'files': [{'name': str, 'content': str}], 'error': str}
    """
    try:
        import tkinter as tk
        from tkinter import filedialog
        from docx import Document

        # 파일 선택 대화상자 (다중 선택)
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        file_paths = filedialog.askopenfilenames(
            title='파일 선택 (여러 개 선택 가능)',
            filetypes=[
                ('텍스트/Word 파일', '*.txt;*.docx'),
                ('텍스트 파일', '*.txt'),
                ('Word 문서', '*.docx'),
                ('모든 파일', '*.*')
            ]
        )

        root.destroy()

        if not file_paths:
            return {'success': False, 'error': '파일 선택이 취소되었습니다.', 'cancelled': True}

        files = []
        errors = []

        for file_path in file_paths:
            filename = os.path.basename(file_path)

            try:
                if file_path.lower().endswith('.docx'):
                    doc = Document(file_path)
                    paragraphs = []
                    for para in doc.paragraphs:
                        text = para.text
                        if text.strip():
                            paragraphs.append(text)
                    content = '\n'.join(paragraphs)
                else:
                    # 텍스트 파일
                    encodings = ['utf-8', 'cp949', 'euc-kr', 'utf-16']
                    content = None
                    for enc in encodings:
                        try:
                            with open(file_path, 'r', encoding=enc) as f:
                                content = f.read()
                            break
                        except:
                            continue

                    if content is None:
                        errors.append(f'{filename}: 인코딩 인식 실패')
                        continue

                files.append({
                    'name': filename,
                    'content': content,
                    'size': os.path.getsize(file_path)
                })

            except Exception as e:
                errors.append(f'{filename}: {str(e)}')

        if not files:
            return {'success': False, 'error': '읽을 수 있는 파일이 없습니다.\n' + '\n'.join(errors)}

        return {
            'success': True,
            'files': files,
            'count': len(files),
            'errors': errors if errors else None
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def open_folder_in_explorer(folder_path):
    """폴더를 탐색기에서 엽니다."""
    if os.path.isdir(folder_path):
        if sys.platform == 'win32':
            os.startfile(folder_path)
        elif sys.platform == 'darwin':
            subprocess.run(['open', folder_path])
        else:
            subprocess.run(['xdg-open', folder_path])


# ===== 브루최적화 (텍스트 분할) 기능 =====

# 절대 줄 시작에 오면 안 되는 한국어 의존명사
DEPENDENT_NOUNS = [
    '적', '것', '때', '줄', '수', '뿐',
    '만큼', '대로', '듯', '바', '터',
    '지', '채', '체', '양', '나위', '따름',
    '번', '데', '점', '리', '나름'
]

# 종결어미 패턴 - 단독으로 한 줄에 있으면 안 됨
ENDING_PATTERNS = [
    '습니다', '니다', '합니다', '입니다', '됩니다', '집니다',
    '봅니다', '옵니다', '줍니다', '씁니다', '갑니다', '납니다',
    '았습니다', '었습니다', '였습니다', '했습니다',
    '세요', '어요', '아요', '네요', '군요', '나요', '죠',
    '해요', '돼요', '봐요', '줘요', '와요', '가요',
    '했어요', '됐어요', '봤어요', '줬어요', '왔어요', '갔어요',
    '다', '요', '까', '죠', '네', '나', '군',
    '니까', '거든', '는데', '어서', '아서',
    '열었습니다', '닫았습니다', '만들었습니다'
]


def _get_pure_length(text):
    """공백 제외한 순수 글자 수 반환"""
    return len(text.replace(' ', ''))


def _is_dependent_noun_start(text):
    """
    텍스트가 의존명사로 시작하는지 확인
    """
    text = text.strip()
    if not text:
        return False

    for noun in DEPENDENT_NOUNS:
        # "적 있으신가요" 또는 딱 "적"인 경우
        if text.startswith(noun + ' ') or text == noun:
            return True
        # "적이" "적은" 같은 조사 결합도 체크
        if len(text) > len(noun) and text.startswith(noun):
            next_char = text[len(noun)]
            if next_char in '이은는을를의도만':
                return True

    return False


def _is_ending_pattern_only(text):
    """
    줄이 종결어미만으로 이루어져 있는지 확인
    예: "했습니다", "열었습니다", "합니다" 등이 단독으로 있는 경우
    """
    text = text.strip()
    if not text:
        return False

    pure_len = _get_pure_length(text)

    # 종결어미만 있는 짧은 줄 (8자 이하)
    if pure_len <= 8:
        for pattern in ENDING_PATTERNS:
            if text.endswith(pattern):
                # 종결어미 앞에 짧은 접두어만 있는 경우
                prefix = text[:-len(pattern)] if len(text) > len(pattern) else ''
                prefix_len = len(prefix.replace(' ', ''))
                # 접두어가 3자 이하면 종결어미 단독으로 간주
                if prefix_len <= 3:
                    return True

    return False


def _is_valid_break_position(text, break_pos, min_length=10):
    """
    끊김점이 유효한지 검증

    조건:
    1. 끊김점 다음이 의존명사로 시작하면 안 됨
    2. 남은 부분이 종결어미만이고 너무 짧으면 안 됨 (핵심!)
    """
    if break_pos >= len(text):
        return True

    next_part = text[break_pos:].strip()

    if not next_part:
        return True

    # 1. 의존명사로 시작하는지 체크
    if _is_dependent_noun_start(next_part):
        return False

    # 2. 남은 부분이 종결어미만이고 너무 짧은지 체크 ⭐ 핵심!
    next_len = _get_pure_length(next_part)

    # 남은 부분이 최소 길이보다 짧고, 종결어미 패턴인 경우 거부
    if next_len < min_length and next_len > 0:
        # 종결어미 패턴 체크 (짧은 형태들)
        short_endings = [
            '습니다', '니다', '합니다', '됩니다', '입니다',
            '했습니다', '겠습니다', '있습니다', '없습니다',
            '봤습니다', '왔습니다', '갔습니다', '줬습니다',
            '열었습니다', '닫았습니다', '말했습니다', '봤습니다',
            '세요', '하세요', '으세요', '보세요',
            '어요', '아요', '해요', '돼요', '네요', '군요',
            '했어요', '됐어요', '봤어요', '왔어요', '갔어요'
        ]

        for ending in short_endings:
            # 남은 부분이 종결어미로 끝나고 짧으면 거부
            if next_part.rstrip('.?!…').endswith(ending):
                # "보였습니다", "말했습니다" 같은 짧은 종결어미
                if next_len <= 8:
                    return False

    return True


def _find_break_point_in_range(text, target, min_len, max_len, options=None):
    """
    지정된 범위 내에서 최적의 끊김점 찾기

    Args:
        text: 분할할 텍스트
        target: 목표 길이 (이상적)
        min_len: 최소 길이
        max_len: 최대 길이
        options: 분할 옵션 (min_length 등)

    Returns:
        끊김점 위치 (None이면 못 찾음)
    """
    if options is None:
        options = {}

    opt_min_length = options.get('min_length', 10)

    # 패턴 정의
    patterns_1 = [
        '습니다', '니다', '합니다', '됩니다', '입니다',
        '했습니다', '겠습니다', '있습니다', '없습니다',
        '하세요', '으세요',
        '가요?', '나요?', '까요?', '어요?', '아요?',
        '네요', '데요', '군요', '구나', '구요',
        '죠', '지요'
    ]

    patterns_2 = [
        '하는데도', '인데도', '지만', '인데', '거든요',
        '니까', '으면', '면서', '으며',
        '하여', '해서', '아서', '어서',
        '려고', '으려고', '다가'
    ]

    patterns_3 = [
        '에서', '으로', '에게', '한테', '께서',
        '부터', '까지', '마다', '처럼', '같이', '보다'
    ]

    char_count = 0
    candidates = []  # (위치, 우선순위, 목표와의 거리)

    for i in range(len(text)):
        if text[i] != ' ':
            char_count += 1

        # 범위 안에 있을 때만 체크
        if min_len <= char_count <= max_len:

            # 0순위: 쉼표
            if text[i] == ',':
                break_pos = i + 1
                # ⭐ min_length 전달하여 남은 부분 길이 체크
                if _is_valid_break_position(text, break_pos, opt_min_length):
                    distance = abs(char_count - target)
                    candidates.append((break_pos, 0, distance, char_count))

            # 1~3순위: 패턴
            for priority, pattern_list in enumerate([patterns_1, patterns_2, patterns_3], 1):
                for pattern in pattern_list:
                    start_idx = max(0, i - len(pattern) + 1)

                    if text[start_idx:i + 1] == pattern:
                        break_pos = i + 1

                        # 패턴 뒤 문자 확인
                        if break_pos < len(text):
                            if text[break_pos] not in ' ,.?!…':
                                continue

                        # ⭐ 의존명사 + 남은 부분 길이 체크
                        if _is_valid_break_position(text, break_pos, opt_min_length):
                            distance = abs(char_count - target)
                            candidates.append((break_pos, priority, distance, char_count))

    # 후보 중 최적 선택: 우선순위 → 목표와의 거리
    if candidates:
        best = min(candidates, key=lambda x: (x[1], x[2]))
        return best[0]

    return None


def _force_split_at_max_length(text, max_length, min_length=10):
    """
    최대 길이에서 강제로 끊기 (공백 찾기)

    Args:
        text: 분할할 텍스트
        max_length: 최대 글자수 (공백 제외)
        min_length: 최소 글자수 (남은 부분 체크용)

    Returns:
        끊김점 위치
    """
    char_count = 0
    target_idx = 0

    # 최대 길이 지점 찾기
    for i in range(len(text)):
        if text[i] != ' ':
            char_count += 1

        if char_count >= max_length:
            target_idx = i + 1
            break

    if target_idx == 0:
        target_idx = len(text)

    # 목표 위치 이전의 가장 가까운 공백
    break_pos = text.rfind(' ', 0, target_idx)

    if break_pos > 0:
        # 의존명사 + 남은 부분 길이 체크
        if _is_valid_break_position(text, break_pos + 1, min_length):
            return break_pos + 1
        else:
            # 유효하지 않으면 더 앞에서
            retry = text.rfind(' ', 0, break_pos)
            if retry > 0:
                if _is_valid_break_position(text, retry + 1, min_length):
                    return retry + 1
                # 그래도 안되면 더 앞에서
                retry2 = text.rfind(' ', 0, retry)
                if retry2 > 0:
                    return retry2 + 1

    # 공백 못 찾으면 목표 위치
    return target_idx


def _split_long_sentence(sentence, options):
    """
    긴 문장 분할 - 최대 길이를 절대 넘으면 안 됨!

    우선순위:
    1. 기준_길이(15자) 근처에서 패턴 찾기 (이상적)
    2. 안 되면 최대_길이(20자) 안에서 패턴 찾기 (확장)
    3. 그래도 안 되면 최대_길이에서 강제 분할 (필수)
    """
    punctuation = '.?!…,;:\'\"''""'
    target_length = options.get('target_length', 15)
    search_range = options.get('search_range', 3)
    max_length = options.get('max_length', 20)

    result = []
    remaining = sentence.strip()

    loop_guard = 0
    max_iterations = 50

    while remaining and loop_guard < max_iterations:
        loop_guard += 1

        pure_length = _get_pure_length(remaining)

        # ===== 최대 이하면 그대로 =====
        if pure_length <= max_length:
            result.append(remaining)
            break

        # ===== 1단계: 기준 근처에서 패턴 찾기 =====
        break_point = _find_break_point_in_range(
            remaining,
            target=target_length,
            min_len=target_length - search_range,
            max_len=target_length + search_range,
            options=options  # ⭐ options 전달
        )

        if break_point:
            chunk = remaining[:break_point].strip()
            chunk_len = _get_pure_length(chunk)

            # 기준 범위 내에서 찾았고, 최대도 안 넘으면 OK
            if chunk_len <= max_length and chunk_len > 0:
                # 문장부호 처리
                next_part = remaining[break_point:].strip()
                while next_part and next_part[0] in punctuation:
                    chunk += next_part[0]
                    next_part = next_part[1:].strip()

                result.append(chunk)
                remaining = next_part
                continue

        # ===== 2단계: 최대 길이까지 확장해서 패턴 찾기 =====
        break_point = _find_break_point_in_range(
            remaining,
            target=max_length,
            min_len=target_length - search_range,
            max_len=max_length,
            options=options  # ⭐ options 전달
        )

        if break_point:
            chunk = remaining[:break_point].strip()
            chunk_len = _get_pure_length(chunk)

            if chunk_len <= max_length and chunk_len > 0:
                # 문장부호 처리
                next_part = remaining[break_point:].strip()
                while next_part and next_part[0] in punctuation:
                    chunk += next_part[0]
                    next_part = next_part[1:].strip()

                result.append(chunk)
                remaining = next_part
                continue

        # ===== 3단계: 강제 분할 (최대 길이에서 공백 찾기) =====
        min_length = options.get('min_length', 10)
        break_point = _force_split_at_max_length(remaining, max_length, min_length)

        if break_point <= 0:
            # 정말 최악의 경우: 첫 단어라도
            first_space = remaining.find(' ')
            if first_space > 0:
                break_point = first_space
            else:
                break_point = max_length

        chunk = remaining[:break_point].strip()

        if not chunk:
            # 비어있으면 강제로 첫 단어
            first_space = remaining.find(' ')
            if first_space > 0:
                chunk = remaining[:first_space]
                break_point = first_space
            else:
                chunk = remaining
                break_point = len(remaining)

        # 문장부호 처리
        next_part = remaining[break_point:].strip()
        while next_part and next_part[0] in punctuation:
            chunk += next_part[0]
            next_part = next_part[1:].strip()

        result.append(chunk)
        remaining = next_part

    if loop_guard >= max_iterations:
        if remaining:
            result.append(remaining)

    return result


def _merge_short_lines(lines, options):
    """
    너무 짧은 줄을 앞/뒤와 병합

    규칙:
    - min_length 미만 줄은 병합 시도
    - 의존명사로 시작하는 줄은 반드시 앞 줄과 병합
    - 종결어미만 있는 줄은 무조건 앞 줄과 병합 (최대 초과해도!)
    - 단, 병합 후 max_length 초과하면 안 함 (종결어미/의존명사 제외)
    """
    min_length = options.get('min_length', 10)
    max_length = options.get('max_length', 18)

    if not lines:
        return lines

    result = []
    i = 0

    while i < len(lines):
        current_line = lines[i].strip()
        current_len = _get_pure_length(current_line)

        # 1순위: 종결어미만 있는 줄은 무조건 앞과 병합 (최대 초과해도!)
        if _is_ending_pattern_only(current_line) and result:
            prev_line = result[-1]
            merged = prev_line + ' ' + current_line
            result[-1] = merged
            i += 1
            continue

        # 2순위: 의존명사로 시작하는 줄은 무조건 앞과 병합
        if _is_dependent_noun_start(current_line) and result:
            prev_line = result[-1]
            merged = prev_line + ' ' + current_line
            merged_len = _get_pure_length(merged)

            # 최대 길이 초과해도 의존명사는 병합 (문법적 필수)
            if merged_len <= max_length + 5:  # 약간의 여유 허용
                result[-1] = merged
                i += 1
                continue

        # 너무 짧은 줄인가?
        if current_len < min_length and i < len(lines) - 1:
            next_line = lines[i + 1].strip()
            merged = current_line + ' ' + next_line
            merged_len = _get_pure_length(merged)

            # 병합해도 최대 허용 안 넘으면 병합
            if merged_len <= max_length:
                result.append(merged)
                i += 2  # 두 줄 건너뛰기
                continue

        # 병합 안 하고 그대로
        result.append(current_line)
        i += 1

    # 2차 패스: 마지막 줄이 너무 짧거나 종결어미만 있으면 앞줄과 병합
    if len(result) > 1:
        last_line = result[-1]
        last_len = _get_pure_length(last_line)

        # 종결어미만 있으면 무조건 병합
        if _is_ending_pattern_only(last_line):
            prev_line = result[-2]
            merged = prev_line + ' ' + last_line
            result[-2] = merged
            result.pop()
        elif last_len < min_length:
            prev_line = result[-2]
            merged = prev_line + ' ' + last_line
            merged_len = _get_pure_length(merged)

            if merged_len <= max_length:
                result[-2] = merged
                result.pop()

    # 3차 패스: 의존명사로 시작하는 줄 다시 체크
    final_result = []
    for line in result:
        if _is_dependent_noun_start(line) and final_result:
            prev_line = final_result[-1]
            final_result[-1] = prev_line + ' ' + line
        else:
            final_result.append(line)

    # 4차 패스: 종결어미만 있는 줄 최종 체크
    final_result2 = []
    for line in final_result:
        if _is_ending_pattern_only(line) and final_result2:
            prev_line = final_result2[-1]
            final_result2[-1] = prev_line + ' ' + line
        else:
            final_result2.append(line)

    return final_result2


def _handle_punctuation(lines):
    """
    문장부호 처리: 문장부호만 있거나 시작하는 줄 병합
    """
    punctuation = '.?!…,;:\'\"''""'

    if len(lines) <= 1:
        return lines

    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 문장부호만 있는 줄
        if all(c in punctuation for c in stripped):
            if result:
                result[-1] = result[-1] + stripped
            else:
                result.append(stripped)
        # 문장부호로 시작하는 줄
        elif stripped[0] in punctuation:
            if result:
                # 앞쪽 문장부호를 이전 줄에 붙이기
                idx = 0
                while idx < len(stripped) and stripped[idx] in punctuation:
                    idx += 1
                result[-1] = result[-1] + stripped[:idx]
                if stripped[idx:].strip():
                    result.append(stripped[idx:].strip())
            else:
                result.append(stripped)
        else:
            result.append(stripped)

    return result


def _split_into_sentences(text):
    """
    텍스트를 문장 단위로 분리
    """
    import re

    # 문장 구분자로 분리 (. ? ! …)
    # 구분자를 유지하면서 분리
    sentences = re.split(r'(?<=[.?!…])\s*', text)

    # 빈 문장 제거 및 정리
    result = []
    for s in sentences:
        s = s.strip()
        if s:
            result.append(s)

    return result


def _process_text(text, options):
    """
    전체 처리 흐름
    1단계: 문장 분리
    2단계: 각 문장 분할
    3단계: 짧은 줄 병합
    4단계: 문장부호 처리
    """
    target_length = options.get('target_length', 15)

    # 1단계: 문장 분리
    sentences = _split_into_sentences(text)

    all_lines = []

    # 2단계: 각 문장별로 분할
    for sentence in sentences:
        pure_length = _get_pure_length(sentence)

        if pure_length <= target_length:
            # 기준 이하면 그대로
            all_lines.append(sentence.strip())
        else:
            # 분할 필요
            split_lines = _split_long_sentence(sentence, options)
            all_lines.extend(split_lines)

    # 3단계: 짧은 줄 병합
    all_lines = _merge_short_lines(all_lines, options)

    # 4단계: 문장부호 처리
    all_lines = _handle_punctuation(all_lines)

    return all_lines


@eel.expose
def process_subtitle_split(file_path, options):
    """
    파일을 읽고 브루최적화 실행

    Args:
        file_path: 파일 경로
        options: {target_length, search_range, min_length, max_length}

    Returns:
        dict: {success, result, original_lines, converted_lines, error}
    """
    try:
        # 파일 읽기
        text = ''
        if file_path.lower().endswith('.docx'):
            try:
                from docx import Document
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text = ' '.join(paragraphs)
            except Exception as e:
                return {'success': False, 'error': f'DOCX 파일 읽기 오류: {str(e)}'}
        else:
            # TXT 파일
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='cp949') as f:
                    text = f.read()

        if not text.strip():
            return {'success': False, 'error': '파일이 비어있습니다.'}

        # 텍스트 전처리 - 줄바꿈을 공백으로 변환
        text = ' '.join(text.split())

        # 문장 단위로 분리 (원본 문장 수)
        sentences = _split_into_sentences(text)
        original_line_count = len(sentences)

        # 전체 처리 실행
        all_lines = _process_text(text, options)

        # 결과 생성
        result_text = '\n'.join(all_lines)

        return {
            'success': True,
            'result': result_text,
            'original_lines': original_line_count,
            'converted_lines': len(all_lines)
        }

    except Exception as e:
        print(f"브루최적화 오류: {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def save_subtitle_result(content, original_filename):
    """
    분할 결과를 TXT로 저장

    Args:
        content: 저장할 내용
        original_filename: 원본 파일명

    Returns:
        dict: {success, path, error}
    """
    try:
        import os
        from tkinter import filedialog
        import tkinter as tk

        # 기본 파일명 생성
        base_name = os.path.splitext(os.path.basename(original_filename))[0]
        default_name = f"{base_name}_분할완료.txt"

        # 저장 대화상자
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        file_path = filedialog.asksaveasfilename(
            defaultextension='.txt',
            filetypes=[('텍스트 파일', '*.txt')],
            initialfile=default_name,
            title='분할 결과 저장'
        )

        root.destroy()

        if not file_path:
            return {'success': False, 'error': '취소됨'}

        # 파일 저장
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return {'success': True, 'path': file_path}

    except Exception as e:
        print(f"저장 오류: {e}")
        return {'success': False, 'error': str(e)}


# ========== PDF 도구 ==========

@eel.expose
def pdf_select_files():
    """PDF 파일 선택 대화상자"""
    try:
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        files = filedialog.askopenfilenames(
            title='PDF 파일 선택',
            filetypes=[('PDF 파일', '*.pdf'), ('모든 파일', '*.*')]
        )

        root.destroy()

        if not files:
            return {'success': False, 'files': []}

        file_list = []
        for path in files:
            file_list.append({
                'path': path,
                'name': os.path.basename(path),
                'size': os.path.getsize(path)
            })

        return {'success': True, 'files': file_list}

    except Exception as e:
        print(f"PDF 파일 선택 오류: {e}")
        return {'success': False, 'error': str(e), 'files': []}


@eel.expose
def pdf_select_folder():
    """폴더에서 PDF 파일 가져오기"""
    try:
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        folder = filedialog.askdirectory(title='PDF 파일이 있는 폴더 선택')

        root.destroy()

        if not folder:
            return {'success': False, 'files': []}

        file_list = []
        for filename in os.listdir(folder):
            if filename.lower().endswith('.pdf'):
                path = os.path.join(folder, filename)
                file_list.append({
                    'path': path,
                    'name': filename,
                    'size': os.path.getsize(path)
                })

        # 이름순 정렬
        file_list.sort(key=lambda x: x['name'].lower())

        return {'success': True, 'files': file_list}

    except Exception as e:
        print(f"PDF 폴더 선택 오류: {e}")
        return {'success': False, 'error': str(e), 'files': []}


@eel.expose
def pdf_select_output_folder():
    """출력 폴더 선택"""
    try:
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        folder = filedialog.askdirectory(title='출력 폴더 선택')

        root.destroy()

        if not folder:
            return {'success': False, 'folder': ''}

        return {'success': True, 'folder': folder}

    except Exception as e:
        print(f"출력 폴더 선택 오류: {e}")
        return {'success': False, 'error': str(e), 'folder': ''}


@eel.expose
def pdf_merge_files(file_paths, output_folder, output_name):
    """여러 PDF 파일을 하나로 합치기"""
    try:
        from PyPDF2 import PdfMerger

        if not file_paths:
            return {'success': False, 'error': '합칠 파일이 없습니다.'}

        # 출력 파일명 처리
        if not output_name.lower().endswith('.pdf'):
            output_name += '.pdf'

        output_path = os.path.join(output_folder, output_name)

        # 기존 파일 있으면 번호 붙이기
        base_name = output_name[:-4]
        counter = 1
        while os.path.exists(output_path):
            output_name = f"{base_name}_{counter}.pdf"
            output_path = os.path.join(output_folder, output_name)
            counter += 1

        merger = PdfMerger()

        for path in file_paths:
            try:
                merger.append(path)
                print(f"[PDF] 추가: {os.path.basename(path)}")
            except Exception as e:
                print(f"[PDF] 파일 추가 실패: {path} - {e}")
                merger.close()
                return {'success': False, 'error': f'파일 처리 실패: {os.path.basename(path)}'}

        merger.write(output_path)
        merger.close()

        print(f"[PDF] 합치기 완료: {output_path}")
        return {'success': True, 'output_path': output_path}

    except ImportError:
        return {'success': False, 'error': 'PyPDF2 라이브러리가 설치되지 않았습니다. pip install PyPDF2'}
    except Exception as e:
        print(f"PDF 합치기 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def pdf_extract_text(file_path, output_folder):
    """PDF에서 텍스트 추출"""
    try:
        from PyPDF2 import PdfReader

        if not os.path.exists(file_path):
            return {'success': False, 'error': '파일을 찾을 수 없습니다.'}

        reader = PdfReader(file_path)
        text_content = []

        for i, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_content.append(f"=== 페이지 {i} ===\n{page_text}\n")

        if not text_content:
            return {'success': False, 'error': '추출할 텍스트가 없습니다.'}

        # 출력 파일명
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_name = f"{base_name}.txt"
        output_path = os.path.join(output_folder, output_name)

        # 기존 파일 있으면 번호 붙이기
        counter = 1
        while os.path.exists(output_path):
            output_name = f"{base_name}_{counter}.txt"
            output_path = os.path.join(output_folder, output_name)
            counter += 1

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_content))

        print(f"[PDF] 텍스트 추출 완료: {output_path}")
        return {'success': True, 'output_path': output_path, 'output_name': output_name, 'pages': len(reader.pages)}

    except ImportError:
        return {'success': False, 'error': 'PyPDF2 라이브러리가 설치되지 않았습니다. pip install PyPDF2'}
    except Exception as e:
        print(f"PDF 텍스트 추출 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def export_thumbnails_to_pdf(thumbnail_urls):
    """
    썸네일 이미지들을 A4 PDF로 출력
    - 세로 방향 (Portrait)
    - 썸네일 크기: 240x135px
    - 가로로 배열
    - 썸네일 간격: 15px
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import mm
        from PIL import Image
        import requests
        from io import BytesIO
        import tempfile
        import datetime

        if not thumbnail_urls or len(thumbnail_urls) == 0:
            return {'success': False, 'error': '출력할 썸네일이 없습니다.'}

        # A4 크기 (세로 방향)
        page_width, page_height = A4  # 210mm x 297mm = 595pt x 842pt

        # 썸네일 설정
        thumb_width = 240  # px를 포인트로 (240px ≈ 180pt)
        thumb_height = 135  # px를 포인트로 (135px ≈ 101pt)
        thumb_spacing = 15  # px를 포인트로 (15px ≈ 11pt)

        # 여백 설정 (mm를 포인트로 변환: 1mm ≈ 2.83pt)
        top_margin = 35 * 2.83  # 35mm
        left_margin = 30 * 2.83  # 30mm
        right_margin = 30 * 2.83  # 30mm
        bottom_margin = 30 * 2.83  # 30mm

        # 한 줄에 2개씩 고정
        thumbs_per_row = 2

        # 출력 파일 경로
        desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = os.path.join(desktop, f'썸네일모음_{timestamp}.pdf')

        # 총 페이지 수 계산 (한 페이지당 들어갈 수 있는 썸네일 수 계산)
        available_height = page_height - top_margin - bottom_margin
        rows_per_page = int((available_height + thumb_spacing) / (thumb_height + thumb_spacing))
        thumbs_per_page = thumbs_per_row * rows_per_page
        total_pages = (len(thumbnail_urls) + thumbs_per_page - 1) // thumbs_per_page

        # 가운데 정렬을 위한 시작 X 좌표 계산
        total_width_needed = (thumbs_per_row * thumb_width) + ((thumbs_per_row - 1) * thumb_spacing)
        start_x = (page_width - total_width_needed) / 2

        # PDF 생성
        c = canvas.Canvas(output_path, pagesize=A4)

        current_x = start_x
        current_y = page_height - top_margin - thumb_height
        thumb_count = 0
        current_page = 1

        # 페이지 번호 표시 함수
        def draw_page_number(canvas_obj, page_num, total):
            canvas_obj.setFont("Helvetica", 10)
            page_text = f"{page_num}/{total}"
            # 오른쪽 상단에 페이지 번호 표시
            text_width = canvas_obj.stringWidth(page_text, "Helvetica", 10)
            canvas_obj.drawString(page_width - right_margin - text_width,
                                 page_height - 20, page_text)

        print(f"[PDF] 썸네일 {len(thumbnail_urls)}개를 PDF로 생성 중...")
        total_thumbs = len(thumbnail_urls)

        for i, url in enumerate(thumbnail_urls):
            try:
                # 진행률 업데이트 (다운로드 단계)
                progress_percent = int((i / total_thumbs) * 100)
                try:
                    eel.updatePdfProgress(f'썸네일 다운로드 중... ({i+1}/{total_thumbs})', progress_percent)
                except:
                    pass

                # 썸네일 다운로드
                response = requests.get(url, timeout=10)
                response.raise_for_status()
                img = Image.open(BytesIO(response.content))

                # 임시 파일로 저장 (reportlab는 PIL 객체를 직접 사용 못함)
                with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_file:
                    img.save(tmp_file.name, 'JPEG')
                    tmp_path = tmp_file.name

                # PDF에 이미지 추가
                c.drawImage(tmp_path, current_x, current_y,
                           width=thumb_width, height=thumb_height,
                           preserveAspectRatio=True, mask='auto')

                # 임시 파일 삭제
                os.unlink(tmp_path)

                # 다음 위치 계산
                current_x += thumb_width + thumb_spacing
                thumb_count += 1

                # 줄 바꿈 필요한 경우
                if thumb_count % thumbs_per_row == 0:
                    current_x = start_x
                    current_y -= thumb_height + thumb_spacing

                    # 페이지 넘김 필요한 경우
                    if current_y < bottom_margin:
                        # 현재 페이지에 페이지 번호 표시
                        draw_page_number(c, current_page, total_pages)
                        c.showPage()
                        current_page += 1
                        current_y = page_height - top_margin - thumb_height
                        current_x = start_x

            except Exception as e:
                print(f"[PDF] 썸네일 처리 실패 ({i+1}/{len(thumbnail_urls)}): {e}")
                continue

        # 마지막 페이지에 페이지 번호 표시
        draw_page_number(c, current_page, total_pages)

        # PDF 저장 진행률 표시
        try:
            eel.updatePdfProgress('PDF 파일 저장 중...', 95)
        except:
            pass

        # PDF 저장
        c.save()

        # 완료 진행률 표시
        try:
            eel.updatePdfProgress('PDF 생성 완료!', 100)
        except:
            pass

        print(f"[PDF] 생성 완료: {output_path}")
        return {
            'success': True,
            'output_path': output_path,
            'thumbnail_count': len(thumbnail_urls)
        }

    except ImportError as e:
        missing_lib = str(e).split("'")[1] if "'" in str(e) else "필요한 라이브러리"
        return {'success': False, 'error': f'{missing_lib} 라이브러리가 설치되지 않았습니다. pip install reportlab pillow requests'}
    except Exception as e:
        print(f"PDF 생성 오류: {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


# ========== 스튜디오 파일 선택 함수 ==========

@eel.expose
def select_folder():
    """폴더 선택 대화상자"""
    try:
        import tkinter as tk
        from tkinter import filedialog

        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        folder = filedialog.askdirectory(title='폴더 선택')

        root.destroy()

        if not folder:
            return None

        return {'path': folder}

    except Exception as e:
        print(f"폴더 선택 오류: {e}")
        return None


@eel.expose
def select_file(extensions=None):
    """파일 선택 대화상자

    Args:
        extensions: 허용할 확장자 리스트 (예: ['srt', 'txt'])
    """
    try:
        import tkinter as tk
        from tkinter import filedialog

        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        # 파일 타입 설정
        if extensions:
            filetypes = [(f'{ext.upper()} 파일', f'*.{ext}') for ext in extensions]
            filetypes.append(('모든 파일', '*.*'))
        else:
            filetypes = [('모든 파일', '*.*')]

        file_path = filedialog.askopenfilename(
            title='파일 선택',
            filetypes=filetypes
        )

        root.destroy()

        if not file_path:
            return None

        return {'path': file_path}

    except Exception as e:
        print(f"파일 선택 오류: {e}")
        return None


@eel.expose
def save_file_dialog(extension='txt'):
    """파일 저장 대화상자

    Args:
        extension: 기본 확장자 (예: 'srt', 'txt')
    """
    try:
        import tkinter as tk
        from tkinter import filedialog

        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        file_path = filedialog.asksaveasfilename(
            title='파일 저장',
            defaultextension=f'.{extension}',
            filetypes=[(f'{extension.upper()} 파일', f'*.{extension}'), ('모든 파일', '*.*')]
        )

        root.destroy()

        if not file_path:
            return None

        return {'path': file_path}

    except Exception as e:
        print(f"파일 저장 대화상자 오류: {e}")
        return None


@eel.expose
def read_file(file_path):
    """파일 내용 읽기

    Args:
        file_path: 읽을 파일 경로
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content

    except Exception as e:
        print(f"파일 읽기 오류: {e}")
        return None


@eel.expose
def write_file(file_path, content):
    """파일 쓰기

    Args:
        file_path: 쓸 파일 경로
        content: 파일 내용
    """
    try:
        # 디렉토리가 없으면 생성
        directory = os.path.dirname(file_path)
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return {'success': True}

    except Exception as e:
        print(f"파일 쓰기 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def save_design_preset(preset_name, settings):
    """영상 탭 프리셋 저장

    Args:
        preset_name: 프리셋 이름
        settings: 저장할 설정 딕셔너리
    """
    try:
        from studio_utils import load_defaults, save_defaults

        # 기존 프리셋 로드
        all_presets = load_defaults()

        # design_presets 키가 없으면 생성
        if 'design_presets' not in all_presets:
            all_presets['design_presets'] = {}

        # 프리셋 저장
        all_presets['design_presets'][preset_name] = settings

        # 파일에 저장
        save_defaults(all_presets)

        print(f"[Design] 프리셋 저장 완료: {preset_name}")
        return {'success': True}

    except Exception as e:
        print(f"프리셋 저장 오류: {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


@eel.expose
def delete_design_preset(preset_name):
    """영상 탭 프리셋 삭제

    Args:
        preset_name: 삭제할 프리셋 이름
    """
    try:
        from studio_utils import load_defaults, save_defaults

        # 기존 프리셋 로드
        all_presets = load_defaults()

        # design_presets 키가 있고 해당 프리셋이 있으면 삭제
        if 'design_presets' in all_presets and preset_name in all_presets['design_presets']:
            del all_presets['design_presets'][preset_name]

            # 파일에 저장
            save_defaults(all_presets)

            print(f"[Design] 프리셋 삭제 완료: {preset_name}")
            return {'success': True}
        else:
            return {'success': False, 'error': '프리셋을 찾을 수 없습니다.'}

    except Exception as e:
        print(f"프리셋 삭제 오류: {e}")
        return {'success': False, 'error': str(e)}


@eel.expose
def load_design_presets():
    """영상 탭 프리셋 목록 로드"""
    try:
        from studio_utils import load_defaults

        all_presets = load_defaults()

        if 'design_presets' in all_presets:
            return {'success': True, 'presets': all_presets['design_presets']}
        else:
            return {'success': True, 'presets': {}}

    except Exception as e:
        print(f"프리셋 로드 오류: {e}")
        return {'success': False, 'error': str(e), 'presets': {}}


if __name__ == '__main__':
    print("=== YouTube 구독 채널 검색 ===")
    print("브라우저에서 앱을 실행합니다...")

    # 사용 가능한 포트 찾기
    port = find_free_port(8000)
    print(f"포트 {port}에서 실행합니다...")

    try:
        # Chrome만 사용 - 전체화면으로 시작
        import ctypes
        user32 = ctypes.windll.user32
        screen_width = user32.GetSystemMetrics(0)
        screen_height = user32.GetSystemMetrics(1)
        eel.start('index.html', size=(screen_width, screen_height), position=(0, 0), port=port, mode='chrome', close_callback=on_close)
    except EnvironmentError:
        print("Chrome 브라우저가 필요합니다. Chrome을 설치해주세요.")
        print("https://www.google.com/chrome/")
        input("Enter를 눌러 종료...")
        sys.exit(1)
